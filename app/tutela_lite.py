"""
Generador de borrador de tutela LIGERO (1 sola llamada Gemini, sin rerank).

Optimizaciones de costo:
  • top_k = 4 fichas, sin rerank LLM
  • max_output_tokens = 1300
  • Cache LRU por hash(descripcion+area) → reusa borradores
  • Prompt con prohibición explícita de inventar y obligación de citar

Detección automática de área a partir de palabras clave para no obligar al
cliente a saber clasificar su caso.
"""

from __future__ import annotations

import hashlib
import io
import re
import time
import threading
from collections import OrderedDict
from typing import Optional

# ── Detección de área por keywords (cero costo) ───────────────────────────────

_KW_AREA = {
    "salud":        ["eps", "ips", "medicamento", "cirugía", "medico", "tratamiento",
                     "oncolo", "cancer", "quimio", "radio", "diagnostico", "hospital",
                     "clinica", "sanitas", "sura", "compensar", "salud total"],
    "pensiones":    ["pension", "colpensiones", "afp", "porvenir", "proteccion",
                     "mesada", "vejez", "invalidez", "sustitucion"],
    "laboral":      ["despido", "trabajo", "salario", "prestaciones", "contrato",
                     "laboral", "empleador", "embarazo", "maternidad", "fuero",
                     "lactancia", "discapacidad laboral", "acoso laboral", "renuncia"],
    "accidentes":   ["accidente", "transito", "soat", "choque", "atropello",
                     "vehiculo", "moto", "polizon", "responsabilidad civil"],
    "insolvencia":  ["insolvencia", "deuda", "embargo", "ejecutivo", "acreedor",
                     "ley 1564", "concurso", "quiebra", "reorganizacion"],
    "derechos_fundamentales": ["fotomulta", "comparendo", "infraccion", "cobro coactivo",
                               "minimo vital", "habeas corpus", "libertad", "libertad inmediata",
                               "mora judicial", "debido proceso"],
}


def detectar_area(texto: str) -> Optional[str]:
    t = (texto or "").lower()
    scores = {}
    for area, kws in _KW_AREA.items():
        scores[area] = sum(1 for k in kws if k in t)
    if not scores: return None
    mejor, n = max(scores.items(), key=lambda x: x[1])
    return mejor if n >= 1 else None


# ── Cache LRU en memoria ──────────────────────────────────────────────────────

_CACHE_MAX = 500
_cache: "OrderedDict[str, dict]" = OrderedDict()
_cache_lock = threading.Lock()


def _normalizar_para_cache(desc: str) -> str:
    """Normaliza para que 'Me niega, Sanitas!' y 'me niega sanitas' peguen al mismo cache."""
    t = (desc or "").lower().strip()
    # Quitar puntuación y múltiples espacios
    t = re.sub(r"[^\w\sáéíóúüñ]", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t[:1500]


def _cache_key(desc: str, area: Optional[str], extra: str = "") -> str:
    raw = ((area or "") + "::" + _normalizar_para_cache(desc) + "::" + (extra or ""))
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def cache_get(desc: str, area: Optional[str], extra: str = "") -> Optional[dict]:
    k = _cache_key(desc, area, extra)
    with _cache_lock:
        if k in _cache:
            _cache.move_to_end(k)
            return _cache[k]
    return None


def cache_put(desc: str, area: Optional[str], value: dict, extra: str = "") -> None:
    k = _cache_key(desc, area, extra)
    with _cache_lock:
        _cache[k] = value
        _cache.move_to_end(k)
        while len(_cache) > _CACHE_MAX:
            _cache.popitem(last=False)


# ── Prompt anti-alucinación ───────────────────────────────────────────────────

DISCLAIMER = (
    "─────────────────────────────────────────\n"
    "⚠ AVISO IMPORTANTE — ESTE BORRADOR NO CONSTITUYE ASESORÍA JURÍDICA\n"
    "─────────────────────────────────────────\n"
    "Este documento es una guía orientativa generada por inteligencia\n"
    "artificial a partir de jurisprudencia pública de la Corte Suprema de\n"
    "Justicia de Colombia. Su contenido NO sustituye la asesoría de un\n"
    "abogado titulado.\n\n"
    "Antes de presentar cualquier acción legal, le recomendamos validar\n"
    "este borrador con un abogado. En Galeano Herrera | Abogados podemos\n"
    "evaluar su caso y acompañarlo en todo el proceso.\n"
    "─────────────────────────────────────────\n\n"
)


def _prompt_publico(descripcion: str, contexto: str, datos_cliente: Optional[dict] = None) -> str:
    from datetime import datetime as _dt
    dc = datos_cliente or {}
    nombre = (dc.get("nombre") or dc.get("name") or "").strip()
    cedula = (dc.get("cedula") or "").strip()
    ciudad = (dc.get("ciudad") or "Bogotá D.C.").strip()
    accionado = (dc.get("accionado") or "").strip()
    telefono = (dc.get("phone") or "").strip()
    email = (dc.get("email") or "").strip()
    fecha = _dt.now().strftime("%d de %B de %Y")

    # Reemplazos de meses en español
    MESES = {"January":"enero","February":"febrero","March":"marzo","April":"abril",
             "May":"mayo","June":"junio","July":"julio","August":"agosto",
             "September":"septiembre","October":"octubre","November":"noviembre","December":"diciembre"}
    for en, es in MESES.items():
        fecha = fecha.replace(en, es)

    block_accionante = f"{nombre}" if nombre else "[COMPLETAR: nombre completo del accionante]"
    block_cedula = f"{cedula}" if cedula else "[COMPLETAR: número de cédula]"
    block_accionado = f"{accionado}" if accionado else "[COMPLETAR: nombre exacto y NIT de la entidad accionada]"
    block_tel = f"+{telefono}" if telefono else "[COMPLETAR: teléfono del accionante]"
    block_email = email if email else "[COMPLETAR: correo del accionante]"

    return f"""Eres jurista colombiano y vas a redactar una SIMULACIÓN MODERADA de acción de tutela.
Esta simulación es una GUÍA ORIENTATIVA para que un abogado titulado revise, profundice y radique.
NO es un documento final listo para firma. Deja espacio explícito para que el abogado agregue valor.

DATOS DEL CASO (úsalos literales si te los dieron, o marca [COMPLETAR con tu abogado]):
  Accionante: {block_accionante}
  Cédula:     {block_cedula}
  Ciudad:     {ciudad}
  Accionado:  {block_accionado}
  Teléfono:   {block_tel}
  Correo:     {block_email}
  Fecha:      {fecha}

REGLAS INNEGOCIABLES:
1. Solo cita radicados que aparezcan EXPRESAMENTE en JURISPRUDENCIA RECUPERADA. NUNCA inventes.
2. Longitud OBJETIVO: 500 a 750 palabras. No más. Deja profundidad para el abogado.
3. Numera 4 a 6 hechos basados en la descripción, sin inventar datos.
4. En FUNDAMENTOS JURÍDICOS: cita 2 a 3 radicados con un párrafo corto por radicado.
   NO redactes la argumentación completa — termina ese bloque con la nota:
   "[Su abogado deberá profundizar la argumentación constitucional específica y las subreglas aplicables.]"
5. En PRETENSIONES: dar la estructura mínima, sin redactar estrategia procesal específica.
   Terminar con la nota: "[Estrategia procesal y pretensiones alternas: su abogado debe ajustarlas al caso.]"
6. En MEDIDA PROVISIONAL: mencionarla como procedente cuando aplique, sin redactar la argumentación jurídica.
   Terminar con la nota: "[La argumentación específica de la medida provisional requiere ajuste por su abogado.]"
7. Entrega solo el documento. Sin comentarios de IA ni meta-explicaciones.
8. Redacta en español formal jurídico colombiano, pero accesible.

JURISPRUDENCIA RECUPERADA:
{contexto}

SITUACIÓN RELATADA POR EL CIUDADANO:
{descripcion}

Genera el BORRADOR con esta estructura exacta (Markdown):

## ACCIÓN DE TUTELA — SIMULACIÓN ORIENTATIVA

**Señor Juez Constitucional de {ciudad} (Reparto)**
**Referencia:** Acción de Tutela de {block_accionante} contra {block_accionado}

{block_accionante}, mayor de edad, identificado con cédula de ciudadanía No. {block_cedula}, domiciliado en {ciudad}, respetuosamente comparezco ante su Despacho para interponer ACCIÓN DE TUTELA contra {block_accionado}, por la vulneración de mis derechos fundamentales, con fundamento en los siguientes:

### I. HECHOS
(Numera 4 a 6 hechos claros, basados en la descripción.)

### II. DERECHOS FUNDAMENTALES VULNERADOS
Lista breve (máximo 3-4 derechos concretos).

### III. FUNDAMENTOS JURÍDICOS (base, requiere profundización)
Cita 2 a 3 radicados con un párrafo breve por radicado explicando su relevancia para el caso.

> **[Su abogado deberá profundizar la argumentación constitucional específica y las subreglas aplicables a este caso concreto.]**

### IV. PROCEDENCIA DE LA ACCIÓN
Valida brevemente en 4 bullets: legitimación activa, legitimación pasiva, inmediatez, subsidiariedad.

### V. PRETENSIONES (estructura base)
1. Amparar los derechos fundamentales vulnerados.
2. Ordenar a {block_accionado} la acción específica solicitada.
3. Cuando aplique, conceder medida provisional urgente.

> **[Estrategia procesal y pretensiones alternas: su abogado debe ajustarlas al caso.]**

### VI. MEDIDA PROVISIONAL
Mención de procedencia si hay riesgo inminente.

> **[La argumentación específica de la medida provisional requiere ajuste por su abogado.]**

### VII. PRUEBAS SUGERIDAS
Lista de documentos típicos del caso. El abogado debe validar cuáles aplican.

### VIII. JURAMENTO
Declaro bajo la gravedad del juramento que no he presentado otra acción de tutela por los mismos hechos.

### IX. NOTIFICACIONES
- Accionante: {block_accionante}, C.C. {block_cedula}. Tel: {block_tel}. Correo: {block_email}. Dirección: [COMPLETAR].
- Accionado: {block_accionado}. [COMPLETAR: dirección].
- **Apoderado (sugerido):** Galeano Herrera | Abogados — contacto@galeanoherrera.co

### X. ANEXOS
1. Copia de la cédula del accionante.
2. Demás pruebas documentales.

{ciudad}, {fecha}

_______________________________
{block_accionante}
C.C. {block_cedula}

---

**Fuentes jurisprudenciales citadas (verifica cada una en relatoria.cortesuprema.gov.co):**
(lista cada radicado en bullet)

**⚠ IMPORTANTE:** Esta simulación es orientativa. Un abogado titulado debe revisarla, profundizar la argumentación y ajustar la estrategia procesal antes de radicar.
"""


# ── Generador principal ───────────────────────────────────────────────────────

def _aplicar_landing_template(template: str, descripcion: str,
                              datos_cliente: Optional[dict],
                              contexto: str) -> str:
    """Reemplaza placeholders {{nombre}}, {{cedula}}, {{ciudad}}, {{accionado}},
    {{descripcion}}, {{contexto_juris}} en el template del vertical."""
    dc = datos_cliente or {}
    repl = {
        "nombre":        dc.get("nombre") or dc.get("name") or "[COMPLETAR nombre]",
        "cedula":        dc.get("cedula") or "[COMPLETAR cédula]",
        "ciudad":        dc.get("ciudad") or "Bogotá D.C.",
        "accionado":     dc.get("accionado") or "[COMPLETAR accionado]",
        "phone":         dc.get("phone") or "[COMPLETAR teléfono]",
        "email":         dc.get("email") or "[COMPLETAR correo]",
        "descripcion":   descripcion,
        "contexto_juris": contexto,
    }
    out = template
    for k, v in repl.items():
        out = out.replace("{{" + k + "}}", str(v))
    return out


def generar_borrador(motor, descripcion: str, area: Optional[str] = None,
                     datos_cliente: Optional[dict] = None,
                     usar_cache: bool = True,
                     landing_cfg: Optional[dict] = None) -> dict:
    """
    motor: instancia de MotorRAG (con FAISS listo o BM25 como mínimo).
    Retorna dict: {draft, fichas, area_detectada, cached, tokens_aprox}
    """
    descripcion = (descripcion or "").strip()
    if len(descripcion) < 30:
        return {"error": "Descripción demasiado corta. Cuente al menos: qué pasó, quién es el accionado y desde cuándo.",
                "draft": "", "fichas": []}

    # Si hay landing_cfg con area_focus, ESE manda y es estricto
    if landing_cfg and landing_cfg.get("area_focus"):
        area = landing_cfg["area_focus"]
    elif not area:
        area = detectar_area(descripcion)

    # Cache: incluye datos que cambian el texto (nombre, cedula, accionado, ciudad)
    cache_extra = ""
    if datos_cliente:
        dc = datos_cliente
        cache_extra = "|".join(str(dc.get(k,"")).strip().lower() for k in
                               ("nombre","cedula","accionado","ciudad"))
    if usar_cache:
        hit = cache_get(descripcion, area, extra=cache_extra)
        if hit:
            hit2 = dict(hit); hit2["cached"] = True
            return hit2

    fichas = motor.buscar(descripcion, k=4, filtro_area=area, hibrido=True)
    if not fichas:
        # Reintentar sin filtro de área
        fichas = motor.buscar(descripcion, k=4, hibrido=True)

    if not fichas:
        return {"error": "No se encontraron precedentes en la base. Reformule incluyendo términos del área legal.",
                "draft": "", "fichas": []}

    contexto = motor._construir_contexto(fichas)
    # Si la landing trae prompt_template propio, úsalo. Si no, el genérico.
    if landing_cfg and (landing_cfg.get("prompt_template") or "").strip():
        prompt = _aplicar_landing_template(
            landing_cfg["prompt_template"], descripcion, datos_cliente, contexto)
    else:
        prompt = _prompt_publico(descripcion, contexto, datos_cliente)

    from google.genai import types as genai_types

    # Intento con cascada de modelos: flash-lite primero (cuota propia),
    # luego 2.0-flash. 5 intentos en total con backoff exponencial.
    MODELOS = ["gemini-2.0-flash-lite", "gemini-2.0-flash"]
    texto = ""
    last_err = None
    BACKOFFS = [0, 4, 8, 15, 25]   # segundos entre intentos

    for intento in range(5):
        modelo = MODELOS[intento % len(MODELOS)]
        if BACKOFFS[intento]:
            time.sleep(BACKOFFS[intento])
        try:
            r = motor._client.models.generate_content(
                model=modelo,
                contents=prompt,
                config=genai_types.GenerateContentConfig(
                    temperature=0.2,
                    max_output_tokens=2600,
                ),
            )
            texto = (r.text or "").strip()
            if texto:
                break
        except Exception as e:
            last_err = e
            msg = str(e)
            # 429 / RESOURCE_EXHAUSTED → seguir al siguiente modelo/intento
            if any(tag in msg for tag in ("429", "RESOURCE_EXHAUSTED", "quota", "Quota")):
                continue
            # Modelo no disponible → probar el siguiente
            if "404" in msg or "NOT_FOUND" in msg:
                continue
            # Errores no recuperables
            return {"error": f"Error generando: {e}",
                    "user_message": "Hubo un problema técnico. Intenta de nuevo en unos minutos.",
                    "draft": "", "fichas": []}

    if not texto:
        return {
            "error": "rate_limited",
            "user_message": (
                "Estamos atendiendo a muchos usuarios en este momento. "
                "Espera 1 minuto y vuelve a intentar. "
                "Si el problema persiste, escríbenos por WhatsApp y te atendemos directamente."
            ),
            "draft": "", "fichas": [],
        }

    if not texto:
        return {"error": str(last_err) if last_err else "Sin respuesta del modelo",
                "user_message": "No se pudo generar el borrador. Intenta de nuevo.", "draft": "", "fichas": []}

    draft = DISCLAIMER + texto + (
        "\n\n─────────────────────────────────────────\n"
        "📞 *Galeano Herrera | Abogados*\n"
        "Su caso merece la asesoría de un especialista.\n"
        "Coordinamos una evaluación gratuita por WhatsApp.\n"
        "─────────────────────────────────────────"
    )

    out = {
        "draft": draft,
        "fichas": [{"id": f["id"], "sala": f.get("sala"), "anio": f.get("anio"),
                    "areas": f.get("areas", [])} for f in fichas],
        "area_detectada": area,
        "cached": False,
        "tokens_aprox": len(prompt)//4 + len(texto)//4,
    }
    if usar_cache:
        cache_put(descripcion, area, out, extra=cache_extra)
    return out


# ── Preview parcial (con efecto blur en frontend) ────────────────────────────

def construir_preview(draft: str, palabras_visibles: int = 260) -> dict:
    """Devuelve el draft particionado: 'visible' (texto plano) y 'oculto' (resto enmascarado)."""
    if not draft: return {"visible": "", "oculto_chars": 0}
    palabras = draft.split()
    if len(palabras) <= palabras_visibles:
        return {"visible": draft, "oculto_chars": 0}
    visible = " ".join(palabras[:palabras_visibles])
    oculto  = " ".join(palabras[palabras_visibles:])
    return {"visible": visible, "oculto_chars": len(oculto), "oculto_palabras": len(palabras)-palabras_visibles}


# ── Render a DOCX (sin dependencias pesadas, fallback a TXT) ─────────────────

def _add_watermark(doc, text: str = "BORRADOR — SIMULACIÓN") -> None:
    """Inserta marca de agua diagonal gris claro en el header de cada página."""
    try:
        from docx.oxml.ns import qn
        from lxml import etree
    except Exception:
        return
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        r = p.add_run()
        # Escape XML special chars in text
        safe = (text or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")
        wm_xml = (
            '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office">'
            '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" '
            'path="m@7,l@8,m@5,21600l@6,21600e">'
            '<v:formulas>'
            '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 @1"/>'
            '<v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>'
            '<v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/><v:f eqn="if @0 @4 21600"/>'
            '<v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>'
            '<v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>'
            '</v:formulas>'
            '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="custom" '
            'o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" '
            'o:connectangles="270,180,90,0" textpathok="t"/>'
            '<v:textpath on="t" fitshape="t"/>'
            '</v:shapetype>'
            '<v:shape id="WM" type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;width:580pt;height:95pt;'
            'z-index:-251648512;mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
            'mso-position-vertical:center;mso-position-vertical-relative:margin;rotation:-40" '
            'fillcolor="#d1d5db" stroked="f">'
            '<v:fill opacity=".55"/>'
            f'<v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt;font-weight:bold" string="{safe}"/>'
            '</v:shape>'
            '</w:pict>'
        )
        try:
            r._r.append(etree.fromstring(wm_xml))
        except Exception:
            pass


def _add_header_footer(doc) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for section in doc.sections:
        # Footer marca
        ftr = section.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("Galeano Herrera | Abogados — contacto@galeanoherrera.co · "
                        "Simulación orientativa · No presentar sin revisión profesional")
        fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)


def borrador_a_docx(draft: str, nombre_cliente: str = "") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return draft.encode("utf-8")

    d = Document()
    section = d.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    style = d.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)

    # Header + footer + watermark
    _add_watermark(d, "SIMULACIÓN — GALEANO HERRERA")
    _add_header_footer(d)

    # Encabezado de marca
    h = d.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("GALEANO HERRERA | ABOGADOS")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x00, 0x23, 0x47)
    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Simulación de Acción de Tutela — Documento orientativo").italic = True

    # Disclaimer box ROJO llamativo al inicio
    box = d.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = box.add_run("⚠ ESTE DOCUMENTO ES UNA SIMULACIÓN ORIENTATIVA.\n"
                     "NO CONSTITUYE ASESORÍA JURÍDICA. "
                     "DEBE SER REVISADO Y AJUSTADO POR UN ABOGADO TITULADO ANTES DE RADICARLO.")
    br.bold = True; br.font.size = Pt(10); br.font.color.rgb = RGBColor(0xc8, 0x10, 0x2e)
    # borde en cajón del párrafo
    pPr = box._p.get_or_add_pPr()
    from docx.oxml.ns import qn as _qn
    from lxml import etree as _et
    pbdr_xml = ('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:top w:val="single" w:sz="12" w:space="6" w:color="C8102E"/>'
                '<w:left w:val="single" w:sz="12" w:space="6" w:color="C8102E"/>'
                '<w:bottom w:val="single" w:sz="12" w:space="6" w:color="C8102E"/>'
                '<w:right w:val="single" w:sz="12" w:space="6" w:color="C8102E"/>'
                '</w:pBdr>')
    try: pPr.append(_et.fromstring(pbdr_xml))
    except Exception: pass
    d.add_paragraph()

    for raw_line in draft.splitlines():
        line = raw_line.rstrip()
        if not line:
            d.add_paragraph(); continue
        if line.startswith("## "):
            p = d.add_paragraph(); r = p.add_run(line[3:].strip())
            r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x00, 0x23, 0x47)
        elif line.startswith("### "):
            p = d.add_paragraph(); r = p.add_run(line[4:].strip())
            r.bold = True; r.font.size = Pt(12)
        elif line.startswith("**") and line.endswith("**") and len(line) < 200:
            p = d.add_paragraph(); p.add_run(line.strip("*")).bold = True
        elif line.startswith("> "):
            # blockquote resaltado (las notas "su abogado deberá...")
            p = d.add_paragraph()
            rr = p.add_run(line[2:].strip("* "))
            rr.italic = True; rr.font.color.rgb = RGBColor(0xC5, 0xA0, 0x59)
        elif line.startswith("- ") or line.startswith("* "):
            d.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            d.add_paragraph(line[line.find(".")+1:].strip(), style="List Number")
        elif line.startswith("─") or set(line.strip()) == {"─"}:
            continue
        else:
            d.add_paragraph(line)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
