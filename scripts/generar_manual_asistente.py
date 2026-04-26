#!/usr/bin/env python3
"""Genera MANUAL_ASISTENTE_DESPACHO.docx — manual operativo para el rol asistente."""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent.parent / "docs" / "MANUAL_ASISTENTE_DESPACHO.docx"
OUT.parent.mkdir(exist_ok=True, parents=True)

AZUL = RGBColor(0x00, 0x23, 0x47); ORO = RGBColor(0xC5, 0xA0, 0x59)
VERDE = RGBColor(0x16, 0xa3, 0x4a); ROJO = RGBColor(0xc8, 0x10, 0x2e)
GRIS = RGBColor(0x6b, 0x72, 0x80)

d = Document(); s = d.sections[0]
s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
d.styles["Normal"].font.name = "Calibri"; d.styles["Normal"].font.size = Pt(11)

def H1(t):
    d.add_page_break(); p = d.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
def H2(t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=AZUL
def H3(t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=ORO
def P(t): return d.add_paragraph(t)
def Bullet(t): d.add_paragraph(t, style="List Bullet")
def Numbered(t): d.add_paragraph(t, style="List Number")
def Spacer(): d.add_paragraph()
def Note(label, t, color=ORO):
    p = d.add_paragraph(); r = p.add_run(label+" "); r.bold=True; r.font.color.rgb=color
    p.add_run(t)
def Tabla(headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0]; rr=c.add_run(h); rr.bold=True; rr.font.color.rgb=AZUL; rr.font.size=Pt(10)
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text=""; rr=c.paragraphs[0].add_run(str(v)); rr.font.size=Pt(10)
    Spacer()

# Portada
for _ in range(3): Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("GALEANO HERRERA | ABOGADOS"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=AZUL
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Manual del Asistente del Despacho"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=ORO
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Dependiente judicial junior + Community Manager + Data entry"); r.italic=True; r.font.size=Pt(13)
for _ in range(8): Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Para el rol asistente del despacho"); r.bold=True; r.font.size=Pt(12)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("(Cronograma, permisos, KPIs y plantillas operativas)"); r.italic=True; r.font.color.rgb=GRIS
for _ in range(6): Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"Versión 1.0 · {datetime.now().strftime('%B %Y')}"); r.font.size=Pt(10); r.font.color.rgb=GRIS

# Capítulo 1
H1("1. Tu rol y por qué existe")
P("El Asistente del Despacho es la pieza que mantiene la operación funcionando mientras los abogados titulados se concentran en lo único que solo ellos pueden hacer: cerrar casos, redactar argumentos finales y representar al cliente en audiencias.")
P("Eres un híbrido entre tres roles tradicionales:")
Bullet("Dependiente judicial junior: triage de casos, organización del expediente, primer contacto.")
Bullet("Community Manager: contenido en redes, respuestas a DMs, mantenimiento de presencia digital.")
Bullet("Data entry / curador: alimentación del cerebro RAG con sentencias y doctrina.")
Spacer()
P("Reportería: directo al titular del despacho. Coordinación operativa con cada abogado por agenda.")

H2("Lo que NO debes hacer (jamás)")
Note("⚠ Importante:", "estas líneas rojas protegen al despacho y a ti misma/o.", color=ROJO)
Bullet("Dar asesoría jurídica. Eres asistente, no abogado. Es delito sin tarjeta profesional.")
Bullet("Cobrar honorarios. Solo el titular o el abogado a cargo definen y cobran.")
Bullet("Firmar tutelas o cualquier documento legal.")
Bullet("Tomar decisiones estratégicas del despacho sin consultar al titular.")
Bullet("Acceder a las claves de admin (galeano + password). Tienes tu propio login en /pro.")
Bullet("Borrar abogados, expedientes ni configuración del sistema.")
Bullet("Compartir datos de clientes con terceros.")

# Capítulo 2
H1("2. Tu acceso al sistema")
H2("URLs y credenciales")
Tabla(
    ["Recurso", "URL", "Cómo entras"],
    [
        ["Tu dashboard /pro", "https://gh-jurisprudencia-csj.onrender.com/pro/login", "Email + password que te entregue el admin"],
        ["Landing pública (referencia)", "https://gh-jurisprudencia-csj.onrender.com", "Sin login"],
        ["Manager WhatsApp (Evolution)", "Ver con el admin", "API key del despacho"],
    ],
)

H2("Lo que ves en tu /pro como asistente")
P("Cuando el admin configure tu cuenta con role='assistant', tu dashboard tendrá tabs adicionales:")
Bullet("📋 Mis Leads + Todos los Leads (puedes asignar)")
Bullet("📅 Agenda (ver agendas de todos los abogados)")
Bullet("⚖ Asistente Jurídico (las 4 herramientas RAG, igual que un abogado)")
Bullet("⏰ Horario (configurar disponibilidad si te asignan citas)")
Bullet("🧠 Cerebro RAG (subir PDFs en modo rápido)")
Bullet("🔧 Cuenta")

# Capítulo 3
H1("3. Cronograma operativo (40 horas / semana)")

H2("Diario · 8 horas")
Tabla(
    ["Bloque", "Hora", "Actividad"],
    [
        ["Apertura", "08:00-08:30", "Café, revisar /pro/leads y mensajes WhatsApp pendientes"],
        ["Triage", "08:30-09:30", "Asignar leads nuevos · primer contacto soft a cada uno"],
        ["Community AM", "09:30-10:00", "Publicar contenido del día en redes (TikTok/IG/FB)"],
        ["Soporte abogados", "10:00-12:00", "Lo que cada abogado solicite (papelería, llamadas, agendamiento)"],
        ["Almuerzo", "12:00-14:00", "—"],
        ["Triage 2", "14:00-15:00", "Segunda tanda de leads + responder DMs"],
        ["Carga RAG", "15:00-16:30", "Subir PDFs nuevos al cerebro (modo rápido)"],
        ["Community PM", "16:30-17:00", "Responder DMs + reportería del día"],
        ["Cierre", "17:00-17:30", "Bitácora del día + handoff al titular"],
    ],
)

H2("Semanal")
Bullet("Lunes 08:00 — Planning con titular (30 min). Métricas semana anterior.")
Bullet("Martes-Jueves — Ejecución estándar.")
Bullet("Viernes 16:00 — Cierre semanal. Reporte de leads, casos cerrados, contenido publicado.")

H2("Mensual")
Bullet("Primer lunes — Revisión de tráfico (30 min) → ajustar campañas con admin.")
Bullet("Mitad de mes — Coordinar con admin enriquecimiento IA del lote pendiente de PDFs.")
Bullet("Último viernes — Presentación de KPIs al titular (15 min).")

# Capítulo 4
H1("4. Gestión de leads paso a paso")

H2("Cuando aparece un lead 'verified'")
Numbered("Entra a /pro → tab Mis Leads o Todos los Leads.")
Numbered("Identifica el área (salud, laboral, accidentes, etc.).")
Numbered("Si el área coincide con un abogado del equipo, asigna a ese abogado.")
Numbered("Si no, asigna al abogado por defecto (marcado con ⭐ en /admin).")
Numbered("Notifica al abogado por WhatsApp interno: 'Te asigné un lead de [área] · cliente [nombre] · resumen [1 línea].'")
Numbered("Marca el lead como 'asignado' en notas.")

H2("Plantilla de primer contacto soft (al cliente)")
P("Solo cuando el abogado titular lo autorice. Texto sugerido:")
Note("Hola [nombre],",
    "te confirmo que tu caso quedó asignado al Dr/Dra. [apellido] de Galeano Herrera | Abogados. "
    "Se pondrá en contacto contigo antes de las [hora]. "
    "Mientras tanto, ten listos: [lista de docs según área]. — [tu nombre], asistente del despacho.")

H2("Documentos típicos por área")
Tabla(
    ["Área", "Documentos a pedir al cliente"],
    [
        ["Salud", "CC, orden médica, historia clínica, comunicación de la EPS, PQR previa"],
        ["Pensiones", "CC, semanas cotizadas (BD Colpensiones), respuesta del fondo, documentos del causante (si sustitución)"],
        ["Laboral", "CC, contrato laboral, despido por escrito, certificación de afiliación EPS, evidencia de causal (embarazo, incapacidad, etc.)"],
        ["Accidentes", "CC, póliza SOAT, registro del accidente (croquis), historia clínica, factura hospitalaria, dictamen pérdida de capacidad"],
        ["Insolvencia", "CC, mandamiento de pago, embargo bancario, certificación cuenta nómina"],
        ["Comparendos", "CC, comparendo, comprobante de notificación (o ausencia), licencia de conducción"],
    ],
)

# Capítulo 5
H1("5. Carga del cerebro RAG")

P("El cerebro RAG es la base de jurisprudencia que el motor IA usa para generar simulaciones y responder consultas. Tu trabajo: alimentarlo de manera ordenada.")

H2("Pipeline de carga")
Numbered("Recibe del titular o abogados los PDFs nuevos a cargar.")
Numbered("Entra a /admin → tab 🧠 Cerebro RAG (necesitas que el admin te dé acceso temporal o trabajen juntos).")
Numbered("Arrastra los PDFs a la zona azul. Usa el modo 'Subida rápida (sin IA)'. Cuesta $0.")
Numbered("Verifica que el sistema haya extraído texto correctamente (cada PDF debe tener > 0 chunks).")
Numbered("Reporta al admin la cantidad cargada al final del día.")

H2("Una vez al mes: enriquecimiento")
P("El admin coordina contigo una sesión de 30 min al mes para hacer click en 'Enriquecer 50 con IA'. Esto procesa los PDFs pendientes y extrae sala, radicado, área, temas y tesis automáticamente.")

H2("Aprobación")
Note("Importante:", "tú NO apruebas PDFs en el RAG. Solo el admin aprueba (es lo que hace que entren al motor). Tu trabajo es cargar y reportar — el admin valida calidad.", color=ROJO)

# Capítulo 6
H1("6. Community Management de redes sociales")

H2("Frecuencia mínima")
Tabla(
    ["Plataforma", "Frecuencia", "Formato"],
    [
        ["TikTok", "5-7 / semana", "Video vertical 30-60s"],
        ["Instagram Reels", "5 / semana", "Mismo asset que TikTok"],
        ["Instagram Posts", "3 / semana", "Carrusel infográfico"],
        ["Facebook", "3-5 / semana", "Post + imagen"],
        ["LinkedIn", "2 / semana", "Post analítico"],
    ],
)

H2("Pilares de contenido (rotar 5)")
Tabla(
    ["Pilar", "% del contenido", "Ejemplo"],
    [
        ["Educativo", "40%", "'¿Sabías que…?', 'Cómo funciona la tutela en 60s'"],
        ["Casos representativos", "25%", "'Un cliente nuestro logró X (anonimizado)'"],
        ["Reacción a noticias jurídicas", "15%", "Sentencia hito → reacción en 24h"],
        ["Detrás de cámaras", "10%", "'Un día en el despacho'"],
        ["CTA directo (solo viernes)", "10%", "'Si te despidieron en embarazo, escríbenos'"],
    ],
)

H2("Lo que NO se publica (Estatuto del Abogado)")
Bullet("Promesas de éxito específico (Art. 38 Ley 1123/2007).")
Bullet("Asesoría legal pública en comentarios (mover a DM).")
Bullet("Nombres de jueces, magistrados o partes específicas sin autorización.")
Bullet("Imágenes de stock con balanzas/martillos: reemplazar por fotos reales del equipo o gráficos editoriales.")

H2("Plantilla de respuesta a DM (genérica)")
Note("Hola [nombre], gracias por escribirnos.",
     "Tu caso suena a [área]. Para darte un análisis serio necesito 3 datos: "
     "(1) ¿qué pasó? · (2) ¿desde cuándo? · (3) ¿tienes algún documento? "
     "Cuando me los compartas, te genero una simulación con jurisprudencia real. "
     "Si necesitas representación, te presento al abogado del área. "
     "— Galeano Herrera | Abogados.")

# Capítulo 7
H1("7. KPIs y reportería")

P("Lo que mides cada semana y reportas al titular el viernes:")
Tabla(
    ["KPI", "Meta", "Cómo medir"],
    [
        ["Leads asignados / día", "≥ 90% del total verificado", "Contar /pro con filtros"],
        ["Tiempo asignación promedio", "< 2h hábiles", "Diferencia verified_at - asignación"],
        ["PDFs cargados / semana", "≥ 20", "Cerebro RAG"],
        ["Publicaciones en redes / semana", "≥ 15 (5 plataformas × 3)", "Manual"],
        ["DMs respondidos en < 4h", "≥ 90%", "Manual"],
        ["NPS del titular contigo", "≥ 8", "Mensual, encuesta corta"],
    ],
)

H2("Plantilla de reporte semanal al titular")
Note("Reporte semana [N] —",
     "Leads recibidos: X · asignados: X · contactados: X. "
     "Casos cerrados (con honorarios): X. "
     "PDFs cargados al RAG: X. "
     "Publicaciones: X (top: '...'). "
     "DMs nuevos: X (Y derivados a lead). "
     "Para revisar la próxima semana: [3 puntos].")

# Capítulo 8
H1("8. Situaciones límite y escalación")

H2("Cliente quiere hablar con un abogado YA")
P("Triángulo de respuesta:")
Numbered("Verificar que el lead esté verificado (con OTP).")
Numbered("Llamar al abogado asignado por WhatsApp interno: '¿Puedes en 10 min?'")
Numbered("Si no puede en <2h, ofrecer cita Meet o agendar para el día siguiente.")
Numbered("Nunca prometer 'el abogado te llama en 5 minutos' si no lo confirmaste.")

H2("Cliente ya pasó las 2 horas sin contacto")
Numbered("Escalar al titular por WhatsApp: 'Lead [nombre], área [X], asignado al Dr Y, sin contactar hace 2h.'")
Numbered("Reasignar al abogado disponible más cercano si el titular lo autoriza.")
Numbered("Documentar en notas del lead.")

H2("DM con queja pública")
Numbered("Responder en privado a la persona pidiendo más detalles.")
Numbered("Avisar al titular del despacho dentro de la siguiente hora.")
Numbered("NO responder en público a la queja sin autorización.")
Numbered("Si la persona insiste públicamente, el titular toma el caso (no tú).")

H2("Cliente cuenta detalles muy graves (violencia, riesgo vital)")
Numbered("Tomar el caso muy en serio. No bromear ni minimizar.")
Numbered("Pasarlo de inmediato al titular, sin filtrar.")
Numbered("Ofrecerle al cliente si quiere agendar cita de urgencia.")
Numbered("Si menciona pensamientos de autolesión, recordarle Línea 106 / 192.")

# Capítulo 9
H1("9. Onboarding tu primer mes")

H2("Semana 1 — Inducción")
Bullet("Lectura de este manual completo + skill `assistant-despacho-junior`.")
Bullet("Tour por la plataforma con el admin (1h).")
Bullet("Crear tu cuenta /pro y configurar tu password.")
Bullet("Agendar 30 min con cada abogado del equipo para conocerse.")
Bullet("Sin asignar leads aún — observa cómo lo hace el titular.")

H2("Semana 2 — Sombra")
Bullet("Asignar leads bajo supervisión del titular.")
Bullet("Responder DMs con plantillas (sin improvisar).")
Bullet("Empezar carga de PDFs al RAG.")
Bullet("Observar 1 reunión con cliente (con autorización).")

H2("Semana 3 — Autonomía supervisada")
Bullet("Asignar leads con autonomía. El titular revisa al final del día.")
Bullet("Empezar publicación en redes (3 plataformas).")
Bullet("Asistir a 1 reunión con cliente y tomar minuta.")

H2("Semana 4 — Operación normal")
Bullet("Operación al ritmo normal del cronograma.")
Bullet("Primer reporte completo de KPIs al titular.")
Bullet("Retroalimentación y ajustes.")

# Cierre
Spacer(); Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("— Fin del Manual del Asistente —"); r.italic=True; r.font.color.rgb=GRIS

d.save(OUT)
print(f"OK · {OUT} · {OUT.stat().st_size:,} bytes")
