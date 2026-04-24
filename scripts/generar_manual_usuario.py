#!/usr/bin/env python3
"""
Genera MANUAL_USUARIO.docx — manual completo para abogados y admin.
Lenguaje no técnico, orientado a uso práctico de la plataforma.

Salida: docs/MANUAL_USUARIO.docx
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = Path(__file__).parent.parent / "docs" / "MANUAL_USUARIO.docx"
OUT.parent.mkdir(exist_ok=True, parents=True)

AZUL = RGBColor(0x00, 0x23, 0x47)
ORO  = RGBColor(0xC5, 0xA0, 0x59)
VERDE = RGBColor(0x16, 0xa3, 0x4a)
ROJO = RGBColor(0xc8, 0x10, 0x2e)
GRIS = RGBColor(0x6b, 0x72, 0x80)

d = Document()

# Márgenes
sec = d.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# Estilo base
style = d.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)


def H1(t):
    d.add_page_break()
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
def H2(t):
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = AZUL
def H3(t):
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ORO
def H4(t):
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
def P(t):
    return d.add_paragraph(t)
def Pb(t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; return p
def Bullet(t):
    d.add_paragraph(t, style="List Bullet")
def Numbered(t):
    d.add_paragraph(t, style="List Number")
def Spacer():
    d.add_paragraph()
def Note(label, t, color=ORO):
    p = d.add_paragraph()
    r = p.add_run(label + " "); r.bold = True; r.font.color.rgb = color
    p.add_run(t)
def Code(t):
    p = d.add_paragraph()
    r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(10)
    r.font.color.rgb = GRIS
def Tabla(headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0]
        run = c.add_run(h); run.bold = True; run.font.color.rgb = AZUL; run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, v in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(v)); run.font.size = Pt(10)
    Spacer()


# ═══════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════
for _ in range(3): Spacer()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GALEANO HERRERA | ABOGADOS")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = AZUL

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manual de Usuario")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ORO

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma Legal con IA · Tutelas en jurisprudencia real")
r.font.size = Pt(13); r.italic = True

for _ in range(8): Spacer()

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Para abogados y administrador del despacho"); r.bold = True; r.font.size = Pt(12)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("(Sin tecnicismos · Lenguaje práctico)"); r.italic = True; r.font.color.rgb = GRIS

for _ in range(6): Spacer()

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Versión 1.0 · {datetime.now().strftime('%B %Y')}")
r.font.size = Pt(10); r.font.color.rgb = GRIS
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("https://gh-jurisprudencia-csj.onrender.com")
r.font.size = Pt(10); r.font.color.rgb = AZUL


# ═══════════════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════════════
H1("Índice")
indice = [
    "Parte 1 — Introducción ........................................... 3",
    "  1. Qué es esta plataforma",
    "  2. Para quién está hecha",
    "  3. Qué problema resuelve",
    "Parte 2 — Estructura legal y jurisprudencial ..................... 5",
    "  4. Marco legal de la acción de tutela en Colombia",
    "  5. Cómo trabaja con jurisprudencia",
    "  6. Habeas Data y compliance (Ley 1581/2012)",
    "Parte 3 — Flujo del cliente (lo que el ciudadano ve) .............. 9",
    "  7. La landing pública",
    "  8. La simulación de tutela",
    "  9. La verificación por WhatsApp",
    "  10. La descarga del documento",
    "  11. El agendamiento de cita",
    "Parte 4 — Manual del Abogado ..................................... 14",
    "  12. Cómo ingresar",
    "  13. Tu dashboard",
    "  14. Tab Agenda",
    "  15. Tab Mis Leads",
    "  16. Tab Asistente Jurídico (las 4 herramientas IA)",
    "  17. Tab Horario (configurar disponibilidad)",
    "  18. Workspace por lead (la “sala de cirugía” del caso)",
    "  19. Tab Cuenta",
    "Parte 5 — Manual del Administrador ............................... 25",
    "  20. Dashboard global y embudo",
    "  21. Gestión de abogados",
    "  22. Gestión de leads",
    "  23. Gestión de citas",
    "  24. Sistema y monitoreo",
    "Parte 6 — Buenas prácticas y protocolo ........................... 30",
    "  25. SLA: 2 horas para responder un lead",
    "  26. Cómo NO usar la simulación",
    "  27. Reglas éticas (Estatuto del Abogado)",
    "Parte 7 — Preguntas frecuentes ................................... 32",
    "Anexos ......................................................... 34",
    "  A. Glosario",
    "  B. Plantillas de WhatsApp para clientes",
    "  C. Plantillas de cierre de venta",
    "  D. Contactos y soporte",
]
for line in indice:
    p = d.add_paragraph(line); p.paragraph_format.space_after = Pt(2)


# ═══════════════════════════════════════════════════════════════════════
# PARTE 1 — INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 1 — Introducción")

H2("1. Qué es esta plataforma")
P("Esta es una plataforma digital del despacho Galeano Herrera | Abogados pensada para hacer dos cosas:")
Numbered("Captar clientes nuevos a través de internet (anuncios, redes, búsquedas en Google).")
Numbered("Acelerar la atención y gestión legal de cada caso usando inteligencia artificial entrenada con jurisprudencia real de la Corte Suprema de Justicia de Colombia.")
Spacer()
P("Visto desde fuera, es una página web (https://gh-jurisprudencia-csj.onrender.com) donde un ciudadano cuenta su problema legal, recibe una simulación de su acción de tutela en menos de un minuto, y queda contactado con un abogado del despacho para que lo represente.")
P("Visto desde dentro, es:")
Bullet("Un panel de control para el administrador (ver leads, abogados, métricas).")
Bullet("Un dashboard personal para cada abogado (sus leads, citas, asistente IA, horario).")
Bullet("Un motor de búsqueda jurisprudencial que cruza cada caso contra 625 sentencias reales de la CSJ (Salas Civil, Laboral, Penal y Plena, períodos 2018-2025).")

H2("2. Para quién está hecha")
H3("Para el ciudadano (cliente potencial)")
P("La persona que entra a la página y siente que un derecho suyo está siendo vulnerado: la EPS le niega un medicamento, el banco le embargó la cuenta, lo despidieron en embarazo, etc. La plataforma le da en minutos un borrador de tutela basado en jurisprudencia real, y lo conecta con un abogado del despacho.")

H3("Para el abogado del despacho")
P("La persona titulada que recibe los leads verificados, los califica, los cierra y eventualmente radica las tutelas. Tiene acceso a un asistente jurídico potente que le permite preparar cada caso en una fracción del tiempo habitual.")

H3("Para el administrador del despacho")
P("El responsable de configurar la operación: dar de alta a los abogados, asignar áreas, monitorear el embudo de conversión, analizar métricas para saber dónde se está cayendo la conversión y qué optimizar.")

H2("3. Qué problema resuelve")
P("El mercado legal colombiano tiene tres dolores tradicionales:")
Tabla(
    ["Dolor del cliente", "Cómo lo resuelve la plataforma"],
    [
        ["No sabe si su caso tiene mérito legal", "La simulación le muestra de inmediato qué dicen las sentencias reales de la Corte sobre situaciones similares."],
        ["No sabe cuánto cuesta un abogado", "La página presenta los rangos de honorarios con transparencia (no hay sorpresas)."],
        ["No quiere papeleo ni esperas largas", "Recibe contacto de un abogado en menos de 2 horas hábiles, vía WhatsApp."],
    ],
)
P("Y para el despacho:")
Tabla(
    ["Dolor del despacho", "Cómo lo resuelve la plataforma"],
    [
        ["Captar clientes es caro y artesanal", "Embudo digital automatizado con tracking de conversión."],
        ["Cada nueva tutela toma horas de redacción", "El asistente IA prepara borradores fundamentados en jurisprudencia real en segundos."],
        ["Es difícil saber si un junior está siendo eficiente", "Métricas individuales por abogado: tasa de cierre, tiempo de respuesta, meta diaria."],
    ],
)


# ═══════════════════════════════════════════════════════════════════════
# PARTE 2 — ESTRUCTURA LEGAL Y JURISPRUDENCIAL
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 2 — Estructura legal y jurisprudencial")

H2("4. Marco legal de la acción de tutela en Colombia")
P("La plataforma está diseñada para apoyar la elaboración y radicación de acciones de tutela bajo el siguiente marco normativo:")

H3("Constitución Política, Artículo 86")
P("Toda persona tendrá acción de tutela para reclamar ante los jueces, en todo momento y lugar, mediante un procedimiento preferente y sumario, por sí misma o por quien actúe a su nombre, la protección inmediata de sus derechos constitucionales fundamentales, cuando quiera que éstos resulten vulnerados o amenazados por la acción o la omisión de cualquier autoridad pública.")

H3("Decreto 2591 de 1991")
P("Reglamenta la acción de tutela. Define los cuatro requisitos clásicos de procedencia:")
Tabla(
    ["Requisito", "Qué significa", "Cómo lo trabaja la plataforma"],
    [
        ["Legitimación por activa", "Quien presenta la tutela debe ser el titular del derecho vulnerado o agente oficioso debidamente justificado.", "El borrador se redacta a nombre del cliente con sus datos exactos (nombre, cédula, ciudad)."],
        ["Legitimación por pasiva", "El accionado debe ser quien efectivamente vulneró el derecho.", "Se solicita al cliente identificar la entidad accionada con nombre exacto y NIT."],
        ["Inmediatez", "El hecho debe ser reciente (típicamente menos de 6 meses).", "El wizard pregunta la fecha del hecho. Si pasa de 6 meses, el abogado debe argumentar la mora."],
        ["Subsidiariedad", "La tutela debe ser el único mecanismo idóneo.", "El borrador incluye un apartado donde el abogado debe sustentarlo, citando jurisprudencia."],
    ],
)

H3("Otras normas relevantes según el área")
Bullet("Salud · Ley 1751 de 2015 (Ley Estatutaria de Salud), Sentencia T-760 de 2008.")
Bullet("Pensiones · Ley 100 de 1993, Acto Legislativo 01 de 2005, jurisprudencia de la Sala Laboral.")
Bullet("Laboral · Código Sustantivo del Trabajo, Ley 1010 de 2006 (acoso), Ley 1257 de 2008 (mujer).")
Bullet("Insolvencia · Ley 1564 de 2012 (Código General del Proceso), Decreto 2677 de 2012.")
Bullet("Habeas Data · Ley 1581 de 2012, Decreto 1377 de 2013.")

H2("5. Cómo trabaja con jurisprudencia")

H3("La base de datos: 625 sentencias de la CSJ")
P("La plataforma tiene indexadas 625 sentencias reales y verificables de la Corte Suprema de Justicia de Colombia, distribuidas así:")
Tabla(
    ["Sala", "Tipo de radicado", "Casos típicos"],
    [
        ["Sala de Casación Civil", "STC", "Tutelas contra particulares, EPS, contratos, familia."],
        ["Sala de Casación Laboral", "STL", "Despidos, fueros, pensiones, seguridad social."],
        ["Sala de Casación Penal", "STP", "Libertad, debido proceso penal, medidas privativas."],
        ["Sala Plena", "(varios)", "Asuntos de mayor trascendencia constitucional."],
    ],
)
P("Las sentencias cubren el período 2018-2025. Cada una está identificada con su radicado oficial, verificable en https://relatoria.cortesuprema.gov.co.")

H3("Cómo busca el sistema los precedentes que aplican a tu caso")
P("El motor combina tres técnicas para encontrar las sentencias más relevantes para cada caso:")
Numbered("Búsqueda por palabras clave (BM25): identifica las sentencias que contienen los términos específicos del caso (por ejemplo, “EPS”, “quimioterapia”, “fuero materno”).")
Numbered("Búsqueda semántica con IA: entiende el sentido aunque las palabras sean diferentes. Si el cliente escribe “me niegan el tratamiento”, encuentra sentencias que hablan de “negativa de servicios médicos”.")
Numbered("Re-ranqueo con IA: una vez identificadas las candidatas, una segunda capa de IA las ordena de más a menos relevantes para el caso específico.")
P("El resultado: las 4 a 6 sentencias más pertinentes para tu caso, en menos de 5 segundos.")

H3("Por qué esto importa: anti-alucinación")
Note("⚠ Crítico:", "todas las IA tienen riesgo de inventar radicados que suenan reales pero no existen. La plataforma está diseñada para que esto NUNCA pase: el modelo solo puede citar radicados que aparezcan en la base de datos indexada. Aun así, tu obligación profesional es verificar cada radicado en la relatoría oficial antes de presentar la tutela.", color=ROJO)

H2("6. Habeas Data y compliance (Ley 1581 de 2012)")
P("Cada cliente que entra a la plataforma debe aceptar tres autorizaciones (presentadas en un único checkbox para reducir fricción, pero legalmente diferenciadas):")
Tabla(
    ["Autorización", "Para qué", "Base legal"],
    [
        ["Términos y Condiciones del servicio", "Define que la simulación NO es asesoría jurídica autónoma.", "Estatuto del Consumidor."],
        ["Tratamiento de datos personales", "Permite procesar nombre, cédula, teléfono y email para generar el borrador.", "Ley 1581/2012, Art. 9."],
        ["Contacto comercial", "Permite que un abogado contacte al cliente por WhatsApp, llamada o correo.", "Ley 1581/2012, Art. 4 lit. f."],
    ],
)
P("El sistema almacena las tres autorizaciones por separado en la base de datos, con timestamp e IP. Si en algún momento una autoridad lo solicita, se puede demostrar el consentimiento expreso.")
Note("Tu deber como abogado:", "respetar las finalidades autorizadas. No usar los datos del cliente para fines distintos a su caso. No compartirlos con terceros. Conservarlos solo el tiempo estrictamente necesario.")


# ═══════════════════════════════════════════════════════════════════════
# PARTE 3 — FLUJO DEL CLIENTE
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 3 — Flujo del cliente (lo que ve el ciudadano)")

H2("7. La landing pública")
P("El cliente entra a https://gh-jurisprudencia-csj.onrender.com. Encuentra:")
Bullet("Una barra superior que se mantiene siempre visible con el botón verde “Empezar ahora”.")
Bullet("Un encabezado con el mensaje: “Te están negando un derecho. Aquí tienes cómo recuperarlo.”")
Bullet("Una franja de estadísticas reales (sentencias indexadas, áreas cubiertas, costo $0 de la simulación).")
Bullet("Un bloque de “Por qué confiar” con tres pilares: autoridad, sin costo oculto, datos protegidos.")
Bullet("Un carrusel horizontal con 34 casos representativos por área (salud, pensiones, laboral, etc.). Cada tarjeta es clicable y lleva directo al wizard con un caso de ejemplo precargado.")
Bullet("El wizard de 5 pasos para describir el caso.")
Bullet("Sección “Cómo funciona” con tres pasos visuales.")
Bullet("FAQ con las preguntas más frecuentes (incluye anclaje de precio).")
Bullet("Footer con enlace al acceso de abogados.")

H2("8. La simulación de tutela")
P("Cuando el cliente completa los 5 pasos del wizard y pulsa “Analizar mi caso”, ocurre lo siguiente:")
Numbered("El motor identifica automáticamente el área legal (salud, pensiones, laboral…) si el cliente no la eligió.")
Numbered("Busca en las 625 sentencias las que más se aproximan a su situación.")
Numbered("Genera un borrador de acción de tutela de 500-750 palabras (versión “moderada”) con: encabezado formal, hechos numerados con base en lo que el cliente narró, derechos vulnerados, fundamentos jurídicos con 2-3 radicados reales, sección de procedencia, pretensiones base, mención de medida provisional, lista de pruebas sugeridas, juramento, notificaciones y firma.")
Numbered("El cliente ve la primera parte del borrador en pantalla; el resto está borroso con un candado.")
Numbered("Si quiere desbloquear y descargar el documento Word, debe completar el siguiente paso (verificación).")
Note("Por qué es “moderada”:", "el borrador está intencionalmente diseñado para que el abogado tenga que aportar valor — la argumentación profunda, la estrategia procesal específica y la redacción definitiva de la medida provisional son tareas explícitamente delegadas al abogado mediante notas en el documento.")

H2("9. La verificación por WhatsApp (OTP)")
P("El cliente registra: nombre completo, cédula, celular Colombia (10 dígitos) y correo. Marca el único checkbox de autorizaciones (que cubre Términos, Habeas Data y contacto comercial).")
P("El sistema le envía por WhatsApp un código de 6 dígitos. El cliente lo ingresa y queda verificado.")
Note("Por qué es importante:", "filtra bots y curiosos. Solo personas reales con celular activo en Colombia entran al embudo. Esto eleva enormemente la calidad del lead.")

H2("10. La descarga del documento")
P("Una vez verificado, el cliente descarga su simulación en formato .docx. El documento incluye:")
Bullet("Marca de agua diagonal “SIMULACIÓN — GALEANO HERRERA” en cada página.")
Bullet("Encabezado oficial con el nombre del despacho.")
Bullet("Caja roja en la primera página: “Este documento es una simulación orientativa. No constituye asesoría jurídica. Debe ser revisado y ajustado por un abogado titulado antes de radicarlo.”")
Bullet("Pie de página con datos de contacto del despacho y disclaimer.")
Bullet("Texto del borrador con secciones claras y notas en cursiva dorada donde el abogado debe profundizar.")

H2("11. El agendamiento de cita")
P("Tras la descarga, al cliente se le ofrece agendar una llamada gratuita de 30 minutos con el abogado asignado. La página le muestra los slots disponibles según el horario configurado por ese abogado.")
P("Si el cliente agenda, recibe confirmación por WhatsApp y la cita queda registrada en el dashboard del abogado. Si no agenda en ese momento, se le promete que un abogado lo contactará en menos de 2 horas hábiles.")
Note("Recordatorios automáticos:", "el sistema envía recordatorio por WhatsApp 24 horas antes y 1 hora antes de la cita. La cita se puede cancelar hasta 60 minutos antes.")


# ═══════════════════════════════════════════════════════════════════════
# PARTE 4 — MANUAL DEL ABOGADO
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 4 — Manual del Abogado")

H2("12. Cómo ingresar")
P("El administrador del despacho te dará tres datos:")
Bullet("URL de acceso: https://gh-jurisprudencia-csj.onrender.com/pro/login")
Bullet("Tu correo (te servirá como usuario).")
Bullet("Una contraseña inicial generada automáticamente (12 caracteres).")
P("Al ingresar por primera vez, te recomendamos cambiar la contraseña inmediatamente desde la pestaña Cuenta. La sesión se mantiene activa por 12 horas; después debes volver a entrar.")

H2("13. Tu dashboard")
P("Al entrar a /pro encontrarás:")
H3("Toggle de disponibilidad")
P("Botón en la parte superior. Cuando está en verde (“Disponible para nuevos casos”), recibes leads. Cuando lo apagas (gris), los nuevos leads se asignan al siguiente abogado disponible. Apágalo cuando estés en audiencia, viajando o de vacaciones.")
H3("Tarjeta de meta diaria")
P("Visualiza tu progreso del día con una barra de avance hacia la meta de 10 tutelas cerradas/día (por abogado, según el plan estratégico del despacho). Te da un mensaje motivacional según tu progreso.")
H3("Mini-stats con semáforo")
Tabla(
    ["Métrica", "Significado", "Color"],
    [
        ["Leads por contactar", "Verificados que aún no marcaste como “contactado”.", "Neutro"],
        ["Citas próximas", "Reuniones agendadas para los próximos días.", "Neutro"],
        ["Tasa de cierre", "% de leads contactados que terminan como “cerrado”.", "Verde si ≥25%, ámbar si ≥15%, rojo si <15%"],
        ["Tiempo respuesta promedio", "Cuánto tardas entre que un lead aparece y lo marcas como contactado.", "Verde si <2h, ámbar si 2-6h, rojo si >6h"],
    ],
)
H3("Embudo personal de 7 días")
P("Te muestra cuántos de tus leads están en cada etapa: Verificados → Contactados → Cerrados. Te ayuda a ver dónde se te están cayendo y enfocar tu energía.")

H2("14. Tab Agenda")
P("Vista de calendario semanal con todas tus citas. Cada celda muestra:")
Bullet("Verde: slot libre.")
Bullet("Azul: cita agendada con un cliente (clic para ver detalle).")
Bullet("Rojo: bloqueado por ti (audiencia, almuerzo, etc.).")
Bullet("Gris: pasado o fuera de tu horario configurado.")
P("Acciones disponibles:")
Bullet("Click en un slot libre: lo bloqueas inmediatamente.")
Bullet("Click en un bloqueo: lo quitas.")
Bullet("Click en una cita: ves los datos del cliente, su caso y el link de Meet (si lo agregaste manualmente).")
Bullet("Navegación con flechas para ver semanas adelante o atrás.")

H2("15. Tab Mis Leads")
P("Tabla con todos los leads que se te han asignado. Filtros disponibles por estado:")
Tabla(
    ["Estado", "Significado"],
    [
        ["preview", "Generó simulación pero no completó registro. No tienes que hacer nada."],
        ["pending_otp", "Se registró pero no verificó el código WhatsApp. Tampoco requiere acción."],
        ["verified", "✅ Confirmó identidad. ESTE ES TU MOMENTO. Contáctalo en menos de 2 horas."],
        ["contacted", "Ya iniciaste conversación. Pendiente de cierre."],
        ["closed", "Caso cerrado (vendido o descartado, dejas notas)."],
    ],
)
P("Acciones por lead:")
Bullet("🧠 Trabajar caso: te lleva al workspace específico del lead (ver sección 18).")
Bullet("Ver: muestra el borrador completo del cliente.")
Bullet("Marcar contactado: cuando ya iniciaste conversación.")
Bullet("Cerrar: cuando ya pagó o decides no continuar el caso.")

H2("16. Tab Asistente Jurídico (las 4 herramientas IA)")
P("Esta es la herramienta más poderosa del sistema para ti. Te permite acceder al motor RAG completo (sin filtros) sobre las 625 sentencias indexadas.")

H3("Herramienta 1 — Consulta libre")
P("Para preguntas abiertas. Ejemplo:")
Code("¿Procede tutela contra particulares cuando hay subordinación económica?")
P("Filtros disponibles: por área (salud, pensiones, etc.), por sala (Civil, Laboral, Penal, Plena) y por año. Útil para refinar la búsqueda cuando ya sabes el contexto.")

H3("Herramienta 2 — Análisis de caso (Protocolo Galeano Herrera)")
P("Te entrega el caso en formato de protocolo:")
Bullet("Clasificación del derecho vulnerado y accionado.")
Bullet("Diagnóstico de procedencia (subsidiariedad, fueros, mínimo vital).")
Bullet("Estrategia recomendada con tiempo y probabilidad de éxito.")
Bullet("Precedentes aplicables citados.")
Bullet("Siguiente paso concreto.")
P("Úsala para preparar la primera llamada con el cliente.")

H3("Herramienta 3 — Generar tutela final (versión completa)")
P("A diferencia de la simulación moderada que ve el cliente, esta herramienta genera la versión COMPLETA, sin notas “tu abogado debe profundizar”. Incluye:")
Bullet("Argumentación constitucional plena.")
Bullet("Subreglas jurisprudenciales explicadas.")
Bullet("Pretensiones específicas por caso.")
Bullet("Medida provisional argumentada.")
P("Es el documento que efectivamente vas a revisar y radicar.")

H3("Herramienta 4 — Línea jurisprudencial")
P("Resumen completo sobre un tema:")
Bullet("Tesis dominante de la Corte (citando radicados).")
Bullet("Evolución en los últimos años.")
Bullet("Criterios que conceden el amparo.")
Bullet("Casos en que la Corte NIEGA (anticipar la defensa contraria).")
Bullet("Cita textual de la sentencia más relevante.")
P("Úsala para preparar argumentos sólidos antes de redactar.")

Note("⚠ Verificación obligatoria:", "siempre verifica cada radicado citado en https://relatoria.cortesuprema.gov.co antes de incluirlo en un documento que vayas a presentar. La plataforma minimiza el riesgo de alucinación, pero la responsabilidad profesional final es tuya.", color=ROJO)

H2("17. Tab Horario (configurar tu disponibilidad)")
P("Aquí defines en qué días y franjas atiendes citas con clientes. Es importante porque:")
Bullet("Los clientes solo pueden agendar contigo en los slots que estén dentro de tu horario.")
Bullet("Los recordatorios y notificaciones se ajustan a tu zona horaria (Bogotá).")

H3("Cómo configurarlo")
P("Verás una grilla con los 7 días de la semana. Cada día tiene:")
Bullet("Un toggle Abierto/Cerrado.")
Bullet("Una lista de franjas horarias (ej. 09:00 → 12:00).")
Bullet("Botón “+ Añadir franja” para agregar más bloques (ej. mañana y tarde con descanso al mediodía).")
Bullet("Botón × para quitar una franja.")

H3("Plantillas rápidas")
P("Para configurarte en segundos, usa una plantilla:")
Bullet("Solo mañanas (8-12): L-V de 8 a 12.")
Bullet("Solo tardes (14-18): L-V de 14 a 18.")
Bullet("Jornada L-V (8-17): L-V de 8 a 12 y 14 a 17 (con almuerzo).")
Bullet("L-S (9-18): incluye sábados.")
Bullet("Sin atención: cierra todos los días.")

H3("Toggle “copiar L-V”")
P("Si lo activas, cualquier cambio que hagas al lunes se replica automáticamente a martes-viernes. Útil para definir un horario laboral uniforme rápido.")

H3("Importante")
P("Cuando guardes, los slots cambian INMEDIATAMENTE para los clientes. Si un cliente ya tenía agendada una cita en una franja que ahora cierras, esa cita NO se cancela automáticamente — debes contactarlo y reprogramar manualmente.")

H2("18. Workspace por lead — la “sala de cirugía” del caso")
P("Acceso: desde la tabla de Mis Leads, click en el botón verde 🧠 “Trabajar caso”. Se abre /pro/lead/[id-del-lead].")
P("Esta es la pantalla más poderosa para tu trabajo diario. Tiene dos columnas:")

H3("Columna izquierda")
H4("📝 Caso del cliente (lectura)")
P("Lo que el cliente escribió en la landing, palabra por palabra. Tienes los datos clave (nombre, cédula, área detectada) y las sentencias base que el motor ya identificó.")

H4("✏️ Editar datos del cliente")
P("Botón en la cabecera. Te abre un panel donde puedes corregir nombre, cédula, teléfono, email, área y descripción del caso. Antes de guardar, el sistema te muestra un resumen de los cambios y pide confirmación.")
P("Útil cuando el cliente cuenta el caso por WhatsApp con más precisión que en el formulario.")

H4("✍️ Complementa el caso (opcional)")
P("Caja para agregar:")
Bullet("Accionado exacto (con NIT si lo tienes).")
Bullet("Derecho fundamental específico vulnerado.")
Bullet("Hechos ampliados que sabes y el cliente no puso.")
P("Toda esta información se concatena automáticamente cuando uses las herramientas IA de la derecha. Cuanto más completo, mejor el resultado.")

H4("📄 Borrador final editable")
P("Aquí pegas o escribes el borrador definitivo. Esta versión REEMPLAZA la simulación original — cuando el cliente vuelva a descargar, recibirá esta nueva versión.")
P("Botones: 💾 Guardar borrador final · 📥 Descargar DOCX.")

H3("Columna derecha — Asistente IA (RAG)")
P("Las mismas 4 herramientas de la pestaña Asistente Jurídico, pero pre-pobladas con los datos del cliente + lo que complementaste a la izquierda. No tienes que volver a escribir nada.")
P("Después de cada generación, dos botones:")
Bullet("📥 Usar este texto como borrador final: copia automáticamente al editor de la izquierda.")
Bullet("📋 Copiar al portapapeles: para pegar en otro lado.")

H3("Acciones rápidas en el header")
Bullet("💬 WhatsApp: abre wa.me/[teléfono] del cliente para que arranques la conversación.")
Bullet("Marcar contactado: actualiza el estado.")
Bullet("Cerrar caso: cuando ya cobraste o decides no continuar.")
Bullet("📥 Descargar borrador actual: te baja el .docx en su versión actual.")

H2("19. Tab Cuenta")
P("Ves tu información (nombre, email, WhatsApp, áreas) y tu estado (Disponible o Pausado).")
P("Botón “Actualizar contraseña”: te recomendamos cambiarla la primera vez y cada 90 días. Mínimo 8 caracteres.")


# ═══════════════════════════════════════════════════════════════════════
# PARTE 5 — MANUAL DEL ADMINISTRADOR
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 5 — Manual del Administrador")

H2("20. Cómo ingresar")
P("URL: https://gh-jurisprudencia-csj.onrender.com/admin")
P("El navegador te pedirá usuario y contraseña (Basic Auth). Te las da el desarrollador o el dueño del proyecto.")

H2("21. Dashboard global y embudo")
P("La parte superior te muestra cinco indicadores clave:")
Tabla(
    ["Métrica", "Qué medir"],
    [
        ["Total leads", "Todos los leads históricos."],
        ["Últimas 24h", "Tracción reciente."],
        ["Verificados", "Pasaron el OTP. Calidad del tráfico."],
        ["Contactados", "El equipo respondió. Eficiencia del SLA."],
        ["Cerrados", "Convertidos en clientes pagos."],
    ],
)
P("Debajo, el “embudo de 7 días” te muestra una barra horizontal por cada etapa: vistas de landing → inicia simulación → simulación lista → registro → OTP verificado → descarga → cita agendada.")
Note("Cómo leerlo:", "donde la barra se acorte fuerte respecto a la anterior, ahí está la fuga. Por ejemplo, si “Inicia simulación → Simulación lista” cae al 50%, hay un problema de UX o de velocidad. Si “Verificado → Cita agendada” cae mucho, hay que revisar el flujo de agendamiento o la disponibilidad de los abogados.")

H2("22. Gestión de abogados")
P("Pestaña “⚖ Abogados”. Verás una tabla con todos los abogados activos.")

H3("Crear un nuevo abogado")
P("Click botón verde “+ Nuevo abogado”. Se abre un modal con:")
Bullet("Nombre completo (visible para el cliente).")
Bullet("WhatsApp (formato 573XXXXXXXXX, sin el +).")
Bullet("Email (será su usuario para entrar al sistema).")
Bullet("Password (botón 🎲 para generar una segura de 12 caracteres).")
Bullet("Áreas que cubre (separadas por coma o asterisco * para todas).")
Bullet("Toggle “es default” (recibe leads sin área específica).")
P("Al guardar, aparece otro modal con las credenciales y el botón “📋 Copiar credenciales” que copia un texto formateado listo para enviar por WhatsApp:")
Code("Galeano Herrera | Abogados\nAcceso pro:\n  URL: https://...\n  Email: ...\n  Password: ...")

H3("Editar el horario de un abogado")
P("Click en “⏰ Horario” en la fila del abogado. Se abre el editor visual semanal idéntico al que usa el abogado en su propia cuenta. Útil para configurar el horario inicial cuando das de alta a alguien.")

H3("Reset password")
P("Botón “🔑 Reset”. Te pide la nueva contraseña y la cambia. Avísale al abogado.")

H3("Toggle de disponibilidad")
P("Botón “On/Off” en la columna “Disp”. Si lo apagas, el abogado deja de recibir leads (pero sí puede seguir trabajando con los que ya tenía).")

H3("Eliminar")
P("Botón × rojo. Pide confirmación. Cuidado: las citas que tenía ese abogado quedan huérfanas (la base de datos guarda referencia pero el lawyer_id queda nulo). Reasignar manualmente si es necesario.")

H2("23. Gestión de leads")
P("Pestaña “📋 Leads”. Tabla con todos los leads del sistema (de todos los abogados).")
P("Por cada lead puedes:")
Bullet("Filtrar por estado.")
Bullet("Click en “🧠 Workspace”: te lleva al espacio de trabajo del lead (es el mismo que ve el abogado, pero como admin tienes acceso a todos).")
Bullet("Click en “Borrador”: ve el documento en una ventana nueva.")
Bullet("Marcar contactado / cerrado.")

H2("24. Gestión de citas")
P("Pestaña “📅 Citas”. Tabla con todas las próximas citas del despacho. Te permite ver el panorama completo del trabajo del equipo y detectar cuellos de botella (ej. un abogado con la agenda llena y otros vacíos).")

H2("25. Sistema y monitoreo")
P("Pestaña “🔧 Sistema”. Te muestra el estado de las integraciones:")
Bullet("✅/❌ Gemini API key (motor de IA).")
Bullet("✅/⏳ Índice FAISS (motor de búsqueda jurisprudencial).")
Bullet("✅/❌ UltraMsg WhatsApp (envío de OTPs y notificaciones).")
Bullet("📞 Abogado por defecto configurado.")
Bullet("DEV_MODE on/off (modo desarrollo, no usar en producción real).")
Note("Si algo aparece en rojo:", "contacta al desarrollador. Las variables de entorno se configuran desde el panel de Render (no desde aquí).")


# ═══════════════════════════════════════════════════════════════════════
# PARTE 6 — BUENAS PRÁCTICAS
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 6 — Buenas prácticas y protocolo")

H2("26. SLA: 2 horas para responder un lead")
P("La regla de oro de la conversión legal en Colombia: cada hora que pasa entre la verificación y el primer contacto, la probabilidad de cierre cae aproximadamente un 10%. Por lo tanto:")
Tabla(
    ["Tiempo desde verificación", "Probabilidad de cierre"],
    [
        ["< 5 minutos", "~30%"],
        ["< 1 hora", "~22%"],
        ["< 2 horas hábiles", "~15%"],
        ["< 24 horas", "~8%"],
        ["> 48 horas", "~3%"],
    ],
)
P("El compromiso del despacho es: contactar a TODO lead verificado en menos de 2 horas hábiles (lunes a viernes 8 am – 6 pm).")

H2("27. Cómo NO usar la simulación")
H3("Reglas innegociables para los abogados")
Numbered("Nunca presentes la simulación tal cual al juzgado. Siempre revísala con el asistente, profundiza la argumentación y ajusta a las particularidades del caso.")
Numbered("Nunca cites un radicado que no hayas verificado en https://relatoria.cortesuprema.gov.co.")
Numbered("Nunca le digas al cliente que “la IA garantiza el éxito”. La IA da insumos; tu responsabilidad como abogado titulado es evaluar la viabilidad real.")
Numbered("Nunca uses el sistema con datos hipotéticos para “probar”. Cada generación cuesta dinero y, si el cliente real tiene problemas, los recursos se han ido en pruebas.")
Numbered("Nunca compartas datos de un cliente con terceros sin su autorización expresa.")

H2("28. Reglas éticas (Estatuto del Abogado)")
P("La plataforma es solo una herramienta. Tú sigues siendo el responsable profesional. Recordatorios:")
Bullet("Solo abogado titulado firma y radica las tutelas.")
Bullet("Honorarios anunciados de forma clara antes del cierre.")
Bullet("Conflicto de interés: pre-screening contra base de clientes existentes antes de tomar el caso.")
Bullet("Confidencialidad: las descripciones de casos no se comparten en redes sin consentimiento expreso.")
Bullet("No representas casos donde claramente no hay derecho — eso quema marca y perjudica al cliente.")


# ═══════════════════════════════════════════════════════════════════════
# PARTE 7 — FAQ
# ═══════════════════════════════════════════════════════════════════════
H1("Parte 7 — Preguntas frecuentes")

H2("Para el abogado")

H3("¿Y si la IA me da un radicado equivocado?")
P("El sistema está diseñado para que el modelo solo cite radicados que existan en su base de datos indexada (las 625 sentencias). El riesgo de invención está minimizado. Aun así, tu obligación profesional es verificar cada radicado en la relatoría oficial de la CSJ. Si encuentras alguno dudoso, repórtalo al admin para que se revise.")

H3("¿Cuántas tutelas debo cerrar al día?")
P("La meta del despacho es 10 cerradas/día por abogado activo. La barra de progreso en tu dashboard te ayuda a visualizarlo. Es una meta ambiciosa para mes 6+ de operación; los primeros meses serán menos.")

H3("Si bloqueo un slot, ¿el cliente puede agendarlo igual?")
P("No. El sistema lo descarta automáticamente. Pero si bloqueas un slot que ya tenía cita, esa cita no se cancela — debes contactar al cliente.")

H3("¿Puedo trabajar leads que no son de mi área?")
P("Sí, si tu campo “áreas” incluye “*” (todas) ves todos los leads. Si tienes áreas específicas, ves los leads de esas áreas + los que se te asignen directamente.")

H3("¿Cómo escalo un caso difícil al titular?")
P("Marca el lead como “contactado” con notas detalladas y avisa por WhatsApp al titular. La plataforma permite marcar el caso como “Workspace” compartido para que ambos trabajen.")

H2("Para el administrador")

H3("¿Cuánto cuesta operar la plataforma?")
P("Render web service: $7/mes (plan starter recomendado). UltraMsg: ~$40/mes. Gemini API: <$5/mes para mil leads. Total: ~$50/mes. La inversión real está en la pauta digital (Google/Facebook Ads), no en la infra.")

H3("¿Cómo agrego una nueva área legal?")
P("Las áreas están codificadas en el sistema. Para agregar una nueva (ej. familia), el desarrollador debe modificar el código. Si solo necesitas re-clasificar un caso, edita el lead desde el workspace.")

H3("¿Y si la API de Gemini se queda sin créditos?")
P("El motor lo detecta y muestra al cliente: “Estamos atendiendo a muchos usuarios, intenta en 1 minuto”. Tienes que recargar la cuenta de Google Cloud Platform o esperar al ciclo siguiente.")


# ═══════════════════════════════════════════════════════════════════════
# ANEXOS
# ═══════════════════════════════════════════════════════════════════════
H1("Anexos")

H2("A. Glosario")
Tabla(
    ["Término", "Definición"],
    [
        ["Lead", "Cliente potencial que pasó por la landing y dejó sus datos."],
        ["Lead verificado", "Lead que confirmó su identidad mediante el código WhatsApp (OTP)."],
        ["RAG", "Retrieval-Augmented Generation. Técnica de IA que combina búsqueda de documentos con generación de texto."],
        ["FAISS", "Motor de búsqueda por similitud semántica desarrollado por Meta. La plataforma lo usa para encontrar sentencias parecidas a un caso."],
        ["BM25", "Algoritmo clásico de búsqueda por palabras clave, usado para complementar la búsqueda semántica."],
        ["OTP", "One-Time Password. Código de uso único enviado por WhatsApp para verificar identidad."],
        ["Workspace", "Pantalla del abogado que reúne todos los datos y herramientas para trabajar un caso específico."],
        ["Slot", "Bloque horario de 30 minutos en la agenda de un abogado, libre para que un cliente agende cita."],
        ["Embudo", "Visualización del recorrido del cliente desde que entra a la landing hasta que cierra el caso."],
        ["SLA", "Service Level Agreement. Compromiso de tiempo de respuesta. En este caso: 2 horas hábiles."],
        ["Sticky bar", "Barra que se mantiene visible mientras el usuario hace scroll."],
        ["Habeas Data", "Derecho fundamental al control de los datos personales (Ley 1581 de 2012)."],
    ],
)

H2("B. Plantillas de WhatsApp para el primer contacto")

H3("Salud (EPS niega tratamiento)")
Code("Hola [nombre], soy [Dr/Dra. apellido] de Galeano Herrera | Abogados.\nAcabo de leer tu caso de [EPS] con la negativa del [tratamiento]. Tu situación está expresamente respaldada por sentencias recientes de la Corte Suprema (STC8916-2023, STC6385-2024).\n\nEl proceso típico para casos como el tuyo:\n• Tutela radicada en 24-48 horas\n• Medida provisional puede salir en 48 horas\n• Fallo en 10 días hábiles\n\n¿Te llamo en 15 minutos para revisar los detalles? Necesito que me confirmes si tienes la negativa por escrito.\n\nSaludos.")

H3("Pensiones (Mora Colpensiones)")
Code("Hola [nombre], soy [Dr/Dra. apellido] de Galeano Herrera.\nVi que llevas [X meses] esperando respuesta de Colpensiones. Esto es un caso clásico de mora administrativa que la jurisprudencia respalda.\n\nLa estrategia es presentar tutela por:\n• Derecho de petición (respuesta en 15 días)\n• Mínimo vital (si tu pensión es tu único ingreso)\n\nTiempo estimado del proceso: 10 días para fallo.\n\n¿Conversamos hoy en la tarde? Tengo disponible 4pm o 5pm. ¿Cuál te sirve?")

H3("Laboral (Despido en embarazo)")
Code("Hola [nombre], soy [Dr/Dra. apellido] de Galeano Herrera.\nLamento mucho lo que estás viviendo. Te confirmo que el despido durante embarazo sin permiso del Ministerio del Trabajo está absolutamente prohibido (Sentencia C-005 de 2017).\n\nLas pretensiones típicas de tu caso son:\n• Reintegro inmediato\n• Pago de salarios dejados de percibir desde el despido\n• Indemnización por despido ineficaz\n• Lactancia y vacaciones\n\nTodo esto suma normalmente entre $20-40 millones según tu salario.\n\n¿Tienes 30 minutos esta semana para revisar contigo el caso por Meet? Es totalmente gratis.")

H2("C. Plantillas de cierre de venta")

H3("Cierre por alternativas")
P("Una vez que el cliente tiene claros los pasos:")
Code("[Nombre], puedo empezar mañana tu proceso. ¿Prefieres que firmemos por Google Docs hoy en la tarde o mañana temprano?")

H3("Cierre resumen (objeción silenciosa)")
Code("Entonces tenemos: [resumen del caso en una frase]. Yo te presento la tutela este [día], el juez tiene 10 días para decidir, las medidas provisionales en 48h. Tu inversión es [precio] todo incluido. ¿Seguimos?")

H3("Cierre por riesgo inverso (solo si aplica)")
Code("Si la tutela no pasa admisión por un tema procedimental, te devuelvo el honorario completo. El riesgo lo corremos nosotros.")

H2("D. Contactos y soporte")
P("Soporte técnico de la plataforma: contacta al desarrollador o admin del despacho.")
P("Reportar un bug: describe qué hacías, qué esperabas y qué pasó. Adjunta captura de pantalla si puedes.")
P("Solicitar nueva funcionalidad: documenta el caso de uso (qué problema resuelve y cuántos abogados se beneficiarían).")
Spacer()
P("URL pública: https://gh-jurisprudencia-csj.onrender.com")
P("URL admin: https://gh-jurisprudencia-csj.onrender.com/admin")
P("URL pro (abogados): https://gh-jurisprudencia-csj.onrender.com/pro/login")
P("Documentación técnica completa: docs/DOCUMENTACION.md en el repositorio.")

# Pie de página
Spacer(); Spacer()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Fin del Manual de Usuario —"); r.italic = True; r.font.color.rgb = GRIS
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Galeano Herrera | Abogados · Plataforma Legal con IA"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUL

d.save(OUT)
print(f"OK · {OUT} · {OUT.stat().st_size:,} bytes")
