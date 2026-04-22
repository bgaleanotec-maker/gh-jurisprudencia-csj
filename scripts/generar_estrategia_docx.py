#!/usr/bin/env python3
"""Genera el Word con la estrategia de captación de clientes."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent.parent / "docs" / "ESTRATEGIA_CAPTACION_CLIENTES.docx"
OUT.parent.mkdir(exist_ok=True, parents=True)

AZUL = RGBColor(0x00, 0x23, 0x47)
ORO  = RGBColor(0xC5, 0xA0, 0x59)

d = Document()
sec = d.sections[0]
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(11)


def H1(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(t); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = AZUL
def H2(t):
    p = d.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AZUL
def H3(t):
    p = d.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORO
def P(t, bold=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; return p
def Bullet(t):
    d.add_paragraph(t, style="List Bullet")
def Numbered(t):
    d.add_paragraph(t, style="List Number")
def Spacer():
    d.add_paragraph()
def Tabla(headers, rows, widths=None):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0]
        run = c.add_run(h); run.bold = True; run.font.color.rgb = AZUL
    for r_idx, row in enumerate(rows, 1):
        for c_idx, v in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = str(v)
    return t


# ── Portada ───────────────────────────────────────────────────────────────────
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GALEANO HERRERA | ABOGADOS"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Estrategia de Captación Digital de Clientes"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = ORO
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma: gh-jurisprudencia-csj.onrender.com"); r.italic = True; r.font.size = Pt(11)
Spacer(); Spacer()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Versión 1.0 · Abril 2026 · Documento estratégico interno"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
d.add_page_break()

# ── Resumen ejecutivo ─────────────────────────────────────────────────────────
H1("1. Resumen ejecutivo")
P(
    "Galeano Herrera | Abogados ha desplegado una plataforma digital que convierte el dolor "
    "más recurrente del ciudadano colombiano (negativas de EPS, demoras de Colpensiones, "
    "despidos en embarazo, embargos, accidentes con SOAT) en clientes calificados de la firma. "
    "El sistema combina un motor RAG de jurisprudencia real de la Corte Suprema (625 sentencias "
    "indexadas, 2018–2025) con un embudo de captación verificado por OTP por WhatsApp."
)
H3("Promesa al usuario final")
P(
    "“Tu borrador de tutela en 60 segundos, basado en jurisprudencia real verificable, "
    "y un abogado que te contacta en menos de 2 horas hábiles si lo necesitas.”"
)
H3("Modelo del negocio")
Bullet("El borrador es gratuito (anzuelo de alta calidad).")
Bullet("El producto vendible es la representación: revisión + radicación + seguimiento del fallo.")
Bullet("Tarifas tipo: $50-250k para tutelas; 20-25% sobre condena en laboral; honorarios fijos en pensiones e insolvencia.")
H3("Métricas norte (norte estrella)")
Bullet("Leads verificados / mes (objetivo mes 1: 100; mes 6: 1.500).")
Bullet("Conversión lead verificado → cliente pagante (objetivo: 8-12%).")
Bullet("CAC promedio (objetivo: < $40.000 por cliente).")
Bullet("LTV promedio (objetivo: > $400.000 por cliente).")
d.add_page_break()


# ── Mercado y dolor ───────────────────────────────────────────────────────────
H1("2. Mercado objetivo y dolor del usuario")
H2("2.1 Tamaño del mercado")
P(
    "Colombia presenta más de 600.000 acciones de tutela al año. De ellas, aproximadamente "
    "el 30% corresponde a salud (EPS), 12% a seguridad social (pensiones), 8% a laboral, "
    "y el resto a derechos fundamentales diversos. La barrera de entrada del ciudadano es:"
)
Bullet("Desconocimiento del trámite y de los precedentes que respaldan su caso.")
Bullet("Miedo al costo de un abogado privado (percibido > $1M cuando la realidad puede ser $150k).")
Bullet("Demora del litigio gratuito (Defensoría, consultorios jurídicos universitarios) — semanas o meses.")
Bullet("Ausencia de información clara sobre 'qué tan probable es ganar mi caso'.")

H2("2.2 Persona principal: 'María, 38 años, Bogotá'")
P(
    "Asalariada formal, ingreso $1.5–3.5M, con familia a cargo. Le acaba de pasar algo que la "
    "frustra (la EPS le niega un examen, el banco le embargó la cuenta, su empleador la despidió "
    "estando embarazada). Busca en Google: 'cómo poner una tutela contra Sanitas'. "
    "Llega a la landing. Si en 60 segundos siente que entendemos su caso y le entregamos algo "
    "concreto, deja sus datos. Si no, vuelve a Google y se pierde."
)
H2("2.3 Personas secundarias")
Bullet("Adulto mayor que necesita pensión (caso típico: hijo o nieto navega por él).")
Bullet("Pequeño comerciante con embargo (urgencia alta, decisión rápida).")
Bullet("Empleado informal con accidente de tránsito (SOAT no paga).")

d.add_page_break()


# ── Embudo ────────────────────────────────────────────────────────────────────
H1("3. Embudo de captación")
H2("3.1 Diagrama del embudo")

Tabla(
    ["Etapa", "Acción del usuario", "Conversión esperada", "KPI"],
    [
        ["Visita", "Llega a la landing", "100%", "Sesiones únicas"],
        ["Engagement", "Escribe descripción del caso", "55-65%", "Tasa de inicio del formulario"],
        ["Preview", "Ve borrador parcial generado", "~95%", "Tasa de generación exitosa"],
        ["Registro", "Llena nombre, cédula, celular, correo", "40-55%", "Tasa de registro"],
        ["OTP", "Verifica código por WhatsApp", "85-92%", "Tasa de verificación OTP"],
        ["Descarga", "Descarga DOCX", "97%+", "Tasa de descarga"],
        ["Contacto del abogado", "Abogado escribe en < 2 h", "100% (manual)", "SLA de primera respuesta"],
        ["Calificación", "Responde el primer mensaje", "30-50%", "Tasa de respuesta inicial"],
        ["Cierre", "Paga honorarios o firma poder", "8-15% sobre verificados", "Tasa de cierre"],
    ],
)
Spacer()
P(
    "Con esos números, una landing con 1.000 visitas/mes produce ~12-25 clientes pagantes. "
    "Con tarifa promedio ponderada de $200.000 = $2.4-5M/mes en facturación neta de la firma "
    "antes de costos."
)

H2("3.2 Mejoras de conversión por etapa")
H3("Visita → Engagement (mejorar 55% → 70%)")
Bullet("Landings verticales por dolor específico (eps-niega, despido-embarazo, etc).")
Bullet("Headline en primera persona del cliente: 'Sanitas me niega…' en vez de 'Tutelas de salud'.")
Bullet("Reducir el textarea a 5 líneas visibles para no intimidar.")
H3("Preview → Registro (mejorar 40% → 60%)")
Bullet("Mostrar el preview MÁS visible (palabras visibles 180 → 220) y el blur menos brusco.")
Bullet("Bullet de prueba social arriba del formulario: '+5.000 borradores generados este mes'.")
Bullet("Contador en vivo: 'María en Cali acaba de descargar su borrador'.")
H3("Registro → OTP (mejorar 85% → 95%)")
Bullet("Mensaje WhatsApp con marca clara y código en formato negrita destacado.")
Bullet("Reenvío automático tras 60 seg sin verificar.")
Bullet("Soporte de OTP por SMS como backup (futura iteración).")

d.add_page_break()


# ── Canales ───────────────────────────────────────────────────────────────────
H1("4. Canales de adquisición")
H2("4.1 Mix recomendado mes 1")
Tabla(
    ["Canal", "Inversión mensual", "Leads esperados", "CAC objetivo"],
    [
        ["Google Ads (palabras: 'tutela salud', 'demanda eps')", "$1.500.000", "120-180", "$8.000–12.000"],
        ["Facebook/Instagram Ads (videos del flujo)", "$1.000.000", "80-150", "$7.000–12.500"],
        ["TikTok orgánico (testimonios + casos ganados)", "$300.000", "40-100", "$3.000–7.500"],
        ["SEO local (blogs por área legal)", "$500.000", "30-60", "$8.000–17.000"],
        ["Grupos de WhatsApp (referidos)", "$0 (tiempo)", "10-30", "$0"],
        ["TOTAL", "$3.300.000", "280-520", "$6.300–11.800"],
    ],
)

H2("4.2 Detalle por canal")

H3("Google Ads")
Bullet("Campañas separadas por área (Salud, Pensiones, Laboral). Mejor calidad que mezclar.")
Bullet("Palabras clave de alta intención: 'tutela contra eps', 'me niegan medicamento', 'colpensiones demora pension'.")
Bullet("Excluir: 'gratis' (atrae no calificados), 'consultoría gratuita' (atrae buscadores de pro bono).")
Bullet("Landing por anuncio: NO mandar siempre a / — usar URLs verticales (/eps-niega).")

H3("Facebook & Instagram Ads")
Bullet("Audiencias: 25-55 años, Colombia, intereses en Derecho/Salud/Trabajo.")
Bullet("Creativos: video vertical 9:16 de 30s, casos reales (con permiso) o ilustraciones de la situación.")
Bullet("CTA: 'Generar mi tutela gratis' → directo a landing.")
Bullet("Look-alikes: subir lista de clientes verificados a Meta cada 2 semanas.")

H3("TikTok orgánico")
Bullet("Cuenta de marca: '@galeanoherrera_legal'. Persona detrás: el abogado titular o asociado joven.")
Bullet("Contenido pilares: 1) explicar fallos recientes, 2) reaccionar a casos virales, 3) tips legales rápidos.")
Bullet("Frecuencia: 5 videos/semana. Costo real: 4 horas/semana del abogado + edición.")

H3("SEO")
Bullet("Blog en /blog con artículos largos (>1.500 palabras) por área legal.")
Bullet("Keywords objetivo: ‘qué hacer si la EPS me niega medicamento’, ‘plazo para tutelar despido’.")
Bullet("Linkbuilding: cita en medios (semana.com, eltiempo.com) cuando salgan fallos polémicos.")

H3("Referidos")
Bullet("Después del cierre exitoso: pedir referido + 5 estrellas en Google My Business.")
Bullet("Programa: 10% del honorario del referido como descuento al referente para su próxima necesidad.")

d.add_page_break()


# ── Operación ─────────────────────────────────────────────────────────────────
H1("5. Operación: del lead al cliente")
H2("5.1 SLA de respuesta")
Tabla(
    ["Estado del lead", "SLA primera respuesta", "Canal preferente"],
    [
        ["Verificado (recién OTP)", "≤ 2 h hábiles", "WhatsApp"],
        ["Sin respuesta a primer mensaje", "Día siguiente, antes de 11am", "WhatsApp"],
        ["Sin respuesta tras 72h", "Mensaje de re-engagement", "WhatsApp"],
        ["Sin respuesta tras 7 días", "Último intento útil + cierre", "WhatsApp"],
        ["Sin respuesta tras 14 días", "Archivar", "—"],
    ],
)

H2("5.2 Roles del equipo")
Bullet("Abogado titular: cierra casos complejos, supervisa QC, recibe leads premium (laboral > $20M).")
Bullet("Asistente comercial (1 FTE inicial): primer contacto, calificación, seguimiento, agenda.")
Bullet("Abogado junior (1 FTE inicial): redacción final de tutelas, radicación, atención a fallos.")
Bullet("Marketing externo o interno: mantenimiento de campañas y SEO.")

H2("5.3 Stack tecnológico")
Bullet("Captación: gh-jurisprudencia-csj.onrender.com (Render web service).")
Bullet("CRM ligero: panel admin propio (/admin) — escalar a HubSpot Free cuando se superen 500 leads/mes.")
Bullet("WhatsApp: UltraMsg (~$40/mes) — escalar a WhatsApp Business API oficial cuando se superen 1.000 mensajes/día.")
Bullet("Almacén legal: Google Drive estructurado /Clientes/<año>/<cédula>/.")
Bullet("Facturación electrónica: Siigo o Alegra ($50-100k/mes).")

H2("5.4 KPIs semanales")
Bullet("Leads totales (objetivo creciente).")
Bullet("Leads verificados (% sobre totales).")
Bullet("Leads contactados < 2h (% sobre verificados — meta 90%).")
Bullet("Casos cerrados (cantidad y monto).")
Bullet("CAC y LTV (calculado mensual).")
Bullet("NPS de clientes cerrados (encuesta a los 30 días).")

d.add_page_break()


# ── Compliance ────────────────────────────────────────────────────────────────
H1("6. Cumplimiento legal y reputacional")
H2("6.1 Habeas Data — Ley 1581 de 2012")
Bullet("Política de Tratamiento de Datos publicada y enlazada en cada formulario (ya incluida en /).")
Bullet("Triple checkbox separado: T&C, Habeas Data, autorización comercial. NUNCA preseleccionados.")
Bullet("Botón de revocatoria visible en /derechos-titular (PRÓXIMA ITERACIÓN).")
Bullet("Registro Nacional de Bases de Datos en SIC: trámite obligatorio una vez la base supere 100 mil titulares.")

H2("6.2 Disclaimer de no asesoría jurídica")
Bullet("Banner permanente en el footer de cada borrador descargado.")
Bullet("Disclaimer en el cuerpo del DOCX antes de cualquier contenido.")
Bullet("T&C aclaran que el borrador NO sustituye al abogado — protege contra futuras demandas por mala práctica.")

H2("6.3 Ética profesional (Estatuto del Abogado)")
Bullet("Solo abogados titulados firman las tutelas presentadas. La IA es herramienta, no sustituto.")
Bullet("Honorarios anunciados de forma clara y no engañosa.")
Bullet("Conflicto de interés: pre-screening contra base de clientes existentes antes de tomar caso.")
Bullet("Confidencialidad: descripciones de casos no se comparten en redes sin consentimiento expreso.")

d.add_page_break()


# ── Proyección financiera ─────────────────────────────────────────────────────
H1("7. Proyección financiera 90 días")
H2("7.1 Costos mensuales operativos")
Tabla(
    ["Concepto", "Costo mensual COP"],
    [
        ["Render Web Service (Starter, persistente)", "$30.000"],
        ["UltraMsg (1 instancia, hasta ~10k mensajes)", "$40.000"],
        ["Gemini API (estimado 3.000 borradores/mes)", "$20.000"],
        ["Dominio + email (godaddy/google)", "$25.000"],
        ["Asistente comercial (FTE)", "$2.500.000"],
        ["Abogado junior (FTE)", "$3.500.000"],
        ["Pauta digital", "$3.300.000"],
        ["Contabilidad + Siigo", "$200.000"],
        ["Reserva imprevistos (10%)", "$960.000"],
        ["TOTAL", "$10.575.000"],
    ],
)

H2("7.2 Proyección de ingresos")
Tabla(
    ["Mes", "Leads verificados", "Casos cerrados", "Tarifa promedio", "Ingresos brutos COP"],
    [
        ["Mes 1", "120", "10 (8.3%)", "$200.000", "$2.000.000"],
        ["Mes 2", "300", "30 (10%)", "$220.000", "$6.600.000"],
        ["Mes 3", "550", "60 (10.9%)", "$240.000", "$14.400.000"],
        ["Mes 6", "1.500", "180 (12%)", "$280.000", "$50.400.000"],
    ],
)
P("Punto de equilibrio: mes 3-4. Margen operativo a partir del mes 5.")

H2("7.3 Sensibilidad")
P(
    "Si la tasa de cierre cae a 5% en lugar de 10%, el punto de equilibrio se mueve al mes 6. "
    "Si subimos la inversión publicitaria a $5M/mes y mantenemos 10% de cierre, llegamos al "
    "punto de equilibrio en mes 2 y a $80M de facturación en mes 6."
)

d.add_page_break()


# ── 30/60/90 ──────────────────────────────────────────────────────────────────
H1("8. Plan 30 / 60 / 90 días")
H2("8.1 Días 1-30 (validación)")
Numbered("Activar landing principal y configurar primer abogado en /admin.")
Numbered("Lanzar Google Ads en una sola vertical (recomendado: salud — mayor volumen y dolor).")
Numbered("Crear 5 piezas de contenido orgánico (TikTok + Reels) sobre fallos recientes.")
Numbered("Validar SLA de 2h con primer asistente comercial. Tener ≥80% de cumplimiento.")
Numbered("Cerrar primeros 10 clientes y documentar el ciclo completo (descripción → cierre → fallo).")

H2("8.2 Días 31-60 (escalar)")
Numbered("Agregar 2-3 landings verticales (despido-embarazo, pension-demorada, embargo).")
Numbered("Triplicar inversión publicitaria si el CAC se mantiene < $15.000.")
Numbered("Implementar sistema de testimonios (con autorización del cliente) para landing.")
Numbered("Migrar a Render Starter ($7/mes) + persistent disk para no perder leads en reinicios.")
Numbered("Contratar abogado junior #2 si el pipeline supera 60 casos abiertos.")

H2("8.3 Días 61-90 (consolidar)")
Numbered("Lanzar versión 2.0 con: chatbot de seguimiento post-radicación, recordatorios de fallos.")
Numbered("Programa de referidos formal (10% descuento por cada cierre referido).")
Numbered("Análisis de cohortes: qué áreas tienen mayor LTV → reasignar inversión.")
Numbered("Inscripción en SIC del Registro Nacional de Bases de Datos (si supera 100k titulares).")
Numbered("Evaluar expansión a otras ciudades (Medellín, Cali) con socios locales.")

d.add_page_break()


# ── Anexos ────────────────────────────────────────────────────────────────────
H1("9. Anexos")
H2("9.1 URLs y credenciales del sistema")
Bullet("Landing pública: https://gh-jurisprudencia-csj.onrender.com/")
Bullet("Panel admin: https://gh-jurisprudencia-csj.onrender.com/admin")
Bullet("App pro (abogados): https://gh-jurisprudencia-csj.onrender.com/pro")
Bullet("Repositorio: https://github.com/bgaleanotec-maker/gh-jurisprudencia-csj")
Bullet("Dashboard Render: https://dashboard.render.com/web/srv-d7k432bbc2fs73fruhhg")

H2("9.2 Skills instalables (en /skills)")
Bullet("landing-converter — optimización de conversión.")
Bullet("whatsapp-lead-nurture — secuencias de seguimiento.")
Bullet("qc-tutela — quality control de borradores antes de firmar.")

H2("9.3 Cronograma de implementación de mejoras")
Tabla(
    ["Cuándo", "Mejora", "Impacto esperado"],
    [
        ["Inmediato", "Configurar abogado default en /admin", "Habilitar notificaciones WhatsApp"],
        ["Semana 1", "Render Starter + disk persistente ($30k/mes)", "No perder leads en reinicios"],
        ["Semana 2", "Apagar DEV_MODE en Render", "OTP solo por WhatsApp (seguridad)"],
        ["Semana 2", "Crear /eps-niega y /pension-demorada", "+30% conversión por specificidad"],
        ["Mes 1", "Integrar reCAPTCHA v3", "Bloquear bots sin afectar UX"],
        ["Mes 2", "Email transaccional (SendGrid o Resend)", "Capturar leads sin celular"],
        ["Mes 3", "Webhook a Google Sheets como backup", "Persistencia de leads sin DB"],
    ],
)

# Pie
Spacer()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Documento estratégico interno · Galeano Herrera | Abogados —")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

d.save(OUT)
print(f"OK: {OUT}  ({OUT.stat().st_size} bytes)")
