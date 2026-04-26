#!/usr/bin/env python3
"""Genera MANUAL_JURISPRUDENCIA_COBERTURA.docx — análisis de cobertura del cerebro RAG
con cronograma de curación por área."""

import json, sys
from collections import Counter
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent.parent / "docs" / "MANUAL_JURISPRUDENCIA_COBERTURA.docx"
OUT.parent.mkdir(exist_ok=True, parents=True)

AZUL = RGBColor(0x00, 0x23, 0x47); ORO = RGBColor(0xC5, 0xA0, 0x59)
VERDE = RGBColor(0x16, 0xa3, 0x4a); ROJO = RGBColor(0xc8, 0x10, 0x2e)
GRIS = RGBColor(0x6b, 0x72, 0x80)

# ── Análisis real del JSONL de fichas CSJ ────────────────────────────────────
INDEX_FILE = Path(__file__).parent.parent / "indices" / "fichas_index.jsonl"

por_area = Counter(); por_anio = Counter(); por_sala = Counter(); total = 0
if INDEX_FILE.exists():
    with open(INDEX_FILE, encoding="utf-8") as f:
        for linea in f:
            try:
                fobj = json.loads(linea)
                total += 1
                for a in fobj.get("areas", []): por_area[a] += 1
                por_sala[str(fobj.get("sala","?"))] += 1
                yr = fobj.get("anio")
                if yr: por_anio[str(yr)] += 1
            except Exception:
                continue

# Categorización por cobertura
def coverage_label(n):
    if n >= 100: return "✅ Excelente", VERDE
    if n >= 50:  return "🟢 Buena", VERDE
    if n >= 30:  return "🟡 Aceptable", ORO
    if n >= 10:  return "🟠 Insuficiente", ORO
    return "🔴 Crítica", ROJO

# ── Documento ────────────────────────────────────────────────────────────────
d = Document(); s = d.sections[0]
s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
d.styles["Normal"].font.name = "Calibri"; d.styles["Normal"].font.size = Pt(11)

def H1(t):
    d.add_page_break(); p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=AZUL
def H2(t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=AZUL
def H3(t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=ORO
def P(t): return d.add_paragraph(t)
def Bullet(t): d.add_paragraph(t, style="List Bullet")
def Numbered(t): d.add_paragraph(t, style="List Number")
def Spacer(): d.add_paragraph()
def Tabla(headers, rows):
    t=d.add_table(rows=1+len(rows), cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i].paragraphs[0]; r=c.add_run(h); r.bold=True; r.font.color.rgb=AZUL; r.font.size=Pt(10)
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row):
            c=t.rows[ri].cells[ci]; c.text=""; r=c.paragraphs[0].add_run(str(v)); r.font.size=Pt(10)
    Spacer()

# Portada
for _ in range(3): Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("GALEANO HERRERA | ABOGADOS"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=AZUL
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Manual de Jurisprudencia"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=ORO
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Análisis de cobertura · vacíos · cronograma de curación"); r.italic=True; r.font.size=Pt(13)
for _ in range(8): Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Para el titular del despacho y el asistente"); r.bold=True; r.font.size=Pt(12)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"Versión {datetime.now().strftime('%B %Y')} · {total} fichas analizadas"); r.italic=True; r.font.color.rgb=GRIS

# Cap 1
H1("1. Resumen ejecutivo")
P(f"Al cierre de esta versión del manual, el cerebro RAG cuenta con **{total} fichas** indexadas, "
  "extraídas de los boletines de tutelas de la Sala de Casación Civil, Laboral, Penal y Sala Plena "
  "de la Corte Suprema de Justicia (2018-2025). Esta base sostiene tanto las simulaciones públicas "
  "del cliente como las consultas avanzadas del asistente jurídico para los abogados.")
P("Este manual analiza la cobertura por área legal, identifica vacíos críticos (especialmente en "
  "Salud y Accidentes) y propone un cronograma trimestral de curación a cargo del asistente del despacho.")

H2(f"Distribución por sala (n={total})")
Tabla(["Sala", "Fichas"], [[k, v] for k,v in sorted(por_sala.items(), key=lambda x:-x[1]) if k != "?"])

H2("Distribución por año")
Tabla(["Año", "Fichas"], [[k, v] for k,v in sorted(por_anio.items())])

# Cap 2
H1("2. Cobertura por área legal")
P("Cada ficha puede tener una o varias áreas. La distribución refleja qué tan preparado está el motor para responder consultas reales:")

areas_orden = ["derechos_fundamentales","laboral","pensiones","salud","insolvencia","accidentes"]
rows = []
for a in areas_orden:
    n = por_area.get(a, 0)
    label, _ = coverage_label(n)
    rows.append([a, n, label])
Tabla(["Área", "Fichas", "Cobertura"], rows)

H2("Análisis área por área")

H3(f"Salud · {por_area.get('salud',0)} fichas")
P("Es el área de **mayor demanda real** de tutelas en Colombia (≈ 30% del volumen nacional según Defensoría). "
  "La cobertura actual es insuficiente para soportar una pauta agresiva en esta vertical. "
  "Con la cobertura presente, el motor responde bien a casos genéricos (negativa de medicamento, cirugía no autorizada), "
  "pero falla en subreglas finas (cuidador domiciliario, transporte intermunicipal, salud mental, tratamientos integrales).")
P("Faltan especialmente:")
Bullet("Sentencias hito post-Sentencia C-313/14 (Estatutaria Salud).")
Bullet("Líneas 2020-2024 sobre cuidador domiciliario para pacientes postrados.")
Bullet("Casos de tratamiento integral en enfermedades catastróficas (cáncer, VIH, lupus).")
Bullet("Negativas oncológicas específicas (medicamentos no POS).")
Bullet("Salud mental: jurisprudencia creciente en internación involuntaria, atención psiquiátrica continua.")
Bullet("Transporte médico intermunicipal (paciente rural con tratamiento en otra ciudad).")

H3(f"Accidentes · {por_area.get('accidentes',0)} fichas")
P("Es la cobertura **crítica** del sistema. Casi vacía. Si el despacho pauta esta vertical "
  "(landing /c/accidentes), el motor IA improvisa respuestas sin precedentes — riesgo alto de errores.")
P("Faltan:")
Bullet("Carro fantasma / FOSYGA / ADRES (cuando no hay conductor identificado).")
Bullet("Cobertura SOAT 800 SMDLV agotamiento prematuro.")
Bullet("Indemnización por pérdida de capacidad laboral tras accidente.")
Bullet("Responsabilidad civil del conductor ebrio (líneas Sala Civil).")
Bullet("Atención hospitalaria post-accidente (suspensión por trámite con SOAT).")
Bullet("Incapacidades laborales tras siniestro.")

H3(f"Insolvencia · {por_area.get('insolvencia',0)} fichas")
P("Cobertura básica. Lo más necesario:")
Bullet("Régimen Ley 1564 persona natural no comerciante (jurisprudencia 2022+).")
Bullet("Inembargabilidad de cuenta nómina (líneas modernas C-543/92, T-462/96).")
Bullet("Pensiones inembargables (jurisprudencia 2020+).")
Bullet("Procesos ejecutivos sin notificación válida.")

H3(f"Pensiones · {por_area.get('pensiones',0)} fichas")
P("Cobertura aceptable, pero el área tiene alta demanda y jurisprudencia muy técnica que cambia rápido. "
  "Hay que mantenerla actualizada:")
Bullet("Jurisprudencia post-Acto Legislativo 01/05.")
Bullet("Mora administrativa Colpensiones (T-373/15 y posteriores).")
Bullet("Pensión sustitución compañera permanente (líneas 2022+).")
Bullet("Indemnización sustitutiva (líneas 2022-2024).")
Bullet("Bono pensional (cuando se traslada de AFP a Colpensiones).")

H3(f"Laboral · {por_area.get('laboral',0)} fichas")
P("Buena cobertura. Mantener actualización con tendencia 2025+:")
Bullet("Acoso laboral con perspectiva de género (Ley 1257/08).")
Bullet("Estabilidad laboral reforzada por discapacidad y enfermedad.")
Bullet("Contrato realidad (líneas 2023+, fallos contra entidades públicas).")
Bullet("Despido en periodo de prueba con discapacidad.")

H3(f"Derechos Fundamentales · {por_area.get('derechos_fundamentales',0)} fichas")
P("Cobertura excelente. Es el área 'paraguas' donde caen muchos casos. Mantener el ritmo de carga.")

# Cap 3
H1("3. Cronograma de curación · 90 días")

P("El asistente del despacho dedica ~3 horas semanales a la carga del cerebro RAG. "
  "Con ese ritmo, el siguiente cronograma cierra los vacíos críticos en un trimestre.")

H2("Mes 1 (35 horas del asistente)")
Tabla(["Semana", "Foco", "Meta de fichas", "Fuente"],
[
  ["1", "Salud", "25 sentencias", "Relatoría CSJ + Boletines mensuales"],
  ["2", "Accidentes", "25 sentencias", "Relatoría CSJ + Sala Civil 2020-2025"],
  ["3", "Salud (continuación)", "25 sentencias", "Casos T-760/08 y posteriores"],
  ["4", "Insolvencia + Pensiones", "20 sentencias", "Sala Civil + Sala Laboral"],
])

H2("Mes 2")
Tabla(["Semana", "Foco", "Meta de fichas", "Fuente"],
[
  ["1", "Accidentes (continuación)", "25 sentencias", "Sala Civil + Penal (responsabilidad)"],
  ["2", "Pensiones", "25 sentencias", "Sala Laboral 2020-2025"],
  ["3", "Salud (cierre del vacío)", "50 sentencias", "Es el área con más demanda"],
  ["4", "Auditoría general + dedupe", "Limpieza", "Skill rag-curator"],
])

H2("Mes 3")
Tabla(["Semana", "Foco", "Meta de fichas", "Fuente"],
[
  ["1", "Mantenimiento + tendencias 2025", "20 sentencias", "Boletines recientes"],
  ["2", "Laboral (refresco)", "30 sentencias", "Líneas 2024-2025"],
  ["3", "Doctrina externa", "10 documentos", "Libros, papers, manuales internos"],
  ["4", "Reindex completo FAISS + reporte titular", "Operación", "Admin"],
])

H2("Resultado esperado tras 90 días")
Tabla(["Área", "Hoy", "Tras Mes 3"],
[
  ["Salud", str(por_area.get("salud",0)), "≥ 140 (insuficiente → buena)"],
  ["Accidentes", str(por_area.get("accidentes",0)), "≥ 60 (crítica → aceptable)"],
  ["Insolvencia", str(por_area.get("insolvencia",0)), "≥ 70 (insuficiente → buena)"],
  ["Pensiones", str(por_area.get("pensiones",0)), "≥ 110 (aceptable → excelente)"],
  ["Laboral", str(por_area.get("laboral",0)), "≥ 250 (mantenimiento)"],
  ["Derechos Fundamentales", str(por_area.get("derechos_fundamentales",0)), "Mantener"],
])

# Cap 4
H1("4. Fuentes oficiales para descargar PDFs")

Tabla(["Fuente", "URL", "Qué tiene"],
[
  ["Relatoría CSJ", "https://relatoria.cortesuprema.gov.co", "Sentencias de Sala Civil, Laboral, Penal y Plena. Buscador por radicado, año, sala."],
  ["Relatoría Corte Const.", "https://relatoria.corteconstitucional.gov.co", "Sentencias C, T, SU. Útil para subreglas constitucionales."],
  ["Boletines tutelas CSJ", "https://cortesuprema.gov.co/boletines-anteriores", "Recopilaciones mensuales con extractos curados. Usar cuando es difícil encontrar la sentencia individual."],
  ["DIJ Vigías Constitucionales", "https://www.dij.usta.edu.co", "Doctrina académica. Útil para libros y papers."],
  ["Documentos del despacho", "Drive interno", "Tutelas que ya hemos litigado. Anonimizar antes de cargar al RAG."],
])

# Cap 5
H1("5. Cómo cargar correctamente cada PDF")

H2("Antes de subir")
Numbered("Renombrar el archivo a un formato estándar: STC-XXXX-AAAA.pdf, T-XXX-AA.pdf, C-XXX-AA.pdf. Esto le permite al sistema detectar el radicado automáticamente del nombre.")
Numbered("Verificar que el PDF tenga texto seleccionable (no sea solo imagen escaneada). Si no, no funciona.")
Numbered("Si el PDF tiene >100 páginas, dividir en capítulos antes de subir.")

H2("Al subir")
Numbered("Entrar a /admin → tab 🧠 Cerebro RAG.")
Numbered("Usar 'Subida rápida (sin IA)' para bulks. Cuesta 0 cuota Gemini.")
Numbered("Confirmar que se generen chunks (deben ser >2 por documento normal).")

H2("Una vez al mes")
Numbered("El admin dispara enriquecimiento por lotes (botón '🤖 Enriquecer 50 con IA').")
Numbered("Esto extrae sala, radicado real, área, temas y tesis automáticamente.")
Numbered("El admin revisa y aprueba los que quedan correctos. Los demás los marca como rejected.")
Numbered("Solo los aprobados entran al motor RAG.")

H2("Trimestralmente")
Numbered("Reindexar FAISS para que las búsquedas semánticas incluyan los chunks aprobados nuevos.")
Numbered("El admin lo hace con el botón 'Reindexar todo'.")

# Cap 6
H1("6. Métrica de calidad del cerebro RAG")

H2("Lo que medimos")
Bullet("**Cobertura por área**: cada área debe tener ≥ 50 chunks aprobados.")
Bullet("**Recencia**: 60% de chunks de los últimos 5 años.")
Bullet("**Hit rate del RAG**: ¿qué % de queries de abogados retornan al menos 3 fichas relevantes? Meta: 95%.")
Bullet("**Quejas internas**: si un abogado dice 'el motor no tiene nada de mi tema', es prioridad.")

H2("Plan de auditoría trimestral")
P("Cada 90 días el titular y el asistente revisan juntos:")
Numbered("Las métricas anteriores.")
Numbered("Las queries con 0 resultados (revelan vacíos del cerebro).")
Numbered("Las quejas o sugerencias de los abogados.")
Numbered("Acciones a tomar para el siguiente trimestre.")

# Cap 7
H1("7. Anexos")

H2("A. Glosario rápido")
Tabla(["Término", "Significado"],
[
  ["Ficha CSJ", "Una entrada estructurada extraída de los boletines de tutelas. Tiene radicado, sala, año, áreas, temas y texto."],
  ["Chunk", "Fragmento de un PDF cargado. Cada PDF se divide en chunks de ~1500 caracteres."],
  ["BM25", "Algoritmo de búsqueda por palabras clave. No usa IA. Gratis e ilimitado."],
  ["Embedding", "Representación numérica del significado de un texto. Se genera con Gemini, cuesta cuota."],
  ["FAISS", "Motor de búsqueda por similitud semántica. Combinado con BM25 da el motor RAG completo."],
  ["RAG", "Retrieval-Augmented Generation. Recupera contexto relevante y lo usa para generar respuestas."],
  ["Cobertura", "Cuánta jurisprudencia tenemos por área. Se mide en cantidad de fichas/chunks."],
  ["Curación", "Proceso de revisar, etiquetar, aprobar o rechazar documentos para que el RAG sea de calidad."],
])

H2("B. Skills relacionadas")
Bullet("`rag-architect-engineer` — arquitectura del cerebro y costos.")
Bullet("`rag-curator` — calidad del contenido, dedupe, taxonomía.")
Bullet("`abogado-tutelas-colombia` — marco legal y etapas procesales.")
Bullet("`assistant-despacho-junior` — rol que ejecuta este cronograma.")
Bullet("`jurisprudencia-coverage-audit` — auditoría continua.")

# Cierre
Spacer(); Spacer()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("— Fin del Manual de Jurisprudencia —"); r.italic=True; r.font.color.rgb=GRIS
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Galeano Herrera | Abogados"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=AZUL

d.save(OUT)
print(f"OK · {OUT} · {OUT.stat().st_size:,} bytes")
print(f"  Total fichas analizadas: {total}")
print(f"  Por área: {dict(por_area)}")
